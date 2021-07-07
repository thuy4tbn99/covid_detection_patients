import pandas as pd
import os
import re
from docx import Document
from collections import OrderedDict

#path docx file
file_path = '/Users/user/Downloads/BÁO CÁO FILE WORD/ BC CHUỖI VỰA VE CHAI/BN0000_BN0000_01_LƯƠNG THỊ THANH THUÝ_NGUYỄN THỊ LỆ THUỶ_01072021.docx'
# file_path = '/Users/user/Downloads/BÁO CÁO FILE WORD/BN0000__HƯNG YÊN_TRẦN VĂN TĂNG_030721.docx'
# file_path = '/Users/user/Downloads/BÁO CÁO FILE WORD/BC CHUỖI CHỢ BÌNH ĐIỀN - NHÓM 2+3/BN0000_ LÊ LÂM THỌ_ĐINH THỊ TRIỀU_24062021.docx'
# file_path = '/Users/user/Downloads/covid_path_split_files/arr_path_1.txt'
# file_path = '/Users/user/Downloads/BÁO CÁO FILE WORD/BC 24 QUẬN HUYỆN TỪ 1-7/HUYỆN HÓC MÔN/BN0000_HM_LÊ THỊ HOÀNG_030721.docx'

#create document object
document = Document(file_path)

date_regex = "[0-9]{1,2}/[0-9]{1,2}/[0-9]{4}"
BN_regex = "BN ?\d+"
def extract_Ngay_duong_tinh(paragraph):
    regex = "(Nhận thông tin lúc)"
    regex = re.compile(regex)
    list_match = None
    if regex.search(paragraph.text):
        # list_match = [m for m in regex.findall(paragraph.text)]
        # print(list_match)
        regex = re.compile(date_regex)
        list_match = regex.findall(paragraph.text)
        if len(list_match) == 0:
            regex = re.compile("ngày ?\d{1,2} ?tháng ?\d{1,2} ?năm ?\d{4}")
            list_match = regex.findall(paragraph.text)
            print('hey',list_match)
            if len(list_match) != 0:
                list_match = list_match[0].split()
                info = list_match[1]+'/'+list_match[3]+'/'+list_match[5]
                return info
    return list_match
def extract_Dich_te(paragraph):
    regex = "([Dd]ịch [Tt]ễ)"
    regex = re.compile(regex)
    if regex.search(paragraph.text):
        info = paragraph.text.split(':')
        return info[1].strip()
def extract_Ngay_lay_mau(paragraph):
    regex = "([Dd]ương tính)"
    regex = re.compile(regex)
    if regex.search(paragraph.text):
        regex = re.compile(date_regex)
        list_match = regex.findall(paragraph.text)
        list_match = list(OrderedDict.fromkeys(list_match))
        return list_match
def extract_Tiep_xuc_ca_duong_tinh(paragraph):
    regex = "([Dd]ương tính)"
    regex = re.compile(regex)
    if regex.search(paragraph.text):
        regex = re.compile(BN_regex)
        list_match = regex.findall(paragraph.text)
        list_match = list(OrderedDict.fromkeys(list_match))
        if len(list_match) == 0:
            return ['Chua ro nguon lay']
        else:
            return list_match
    #     p = paragraph.text.lower()
    #     a = re.findall(r'[^\sa-zA-Záàảãạăắằẳẵặâấầẩẫậéèẻẽẹêếềểễệóòỏõọôốồổỗộơớờởỡợíìỉĩịúùủũụưứừửữựýỳỷỹỵđ_]', paragraph.text)
    #     print(a)

def single_patient(document):
    Ngay_lay_mau = []
    Ngay_xet_nghiem_duong_tinh = ''
    Dich_te = ''
    Tiep_xuc_ca_duong_tinh = []
    i = 0
    Thong_tin_ca_benh = []
    for paragraph in document.paragraphs:
        if 'Thông tin ca bệnh' in paragraph.text:
            i += 1
        if i > 0:
            print(paragraph.text)
            Thong_tin_ca_benh.append(paragraph)
        if 'Lịch sử đi lại và tiền sử' in paragraph.text:
            i = 0
            break
    # print(Thong_tin_ca_benh)
    for paragraph in Thong_tin_ca_benh:
        if extract_Ngay_duong_tinh(paragraph) != None:
            # print(extract_Ngay_duong_tinh(paragraph))
            Ngay_xet_nghiem_duong_tinh = extract_Ngay_duong_tinh(paragraph)
        if extract_Dich_te(paragraph) != None:
            # print(extract_Dich_te(paragraph))
            Dich_te = extract_Dich_te(paragraph)
        if extract_Ngay_lay_mau(paragraph) != None:
            # print(extract_Ngay_lay_mau(paragraph))
            Ngay_lay_mau = extract_Ngay_lay_mau(paragraph)
        if extract_Tiep_xuc_ca_duong_tinh(paragraph) != None:
            # print(extract_Tiep_xuc_ca_duong_tinh(paragraph))
            Tiep_xuc_ca_duong_tinh = extract_Tiep_xuc_ca_duong_tinh(paragraph)
    return {'Ngay_xet_nghiem_duong_tinh':Ngay_xet_nghiem_duong_tinh,
              'Dich_te':Dich_te,
              'Ngay_lay_mau':Ngay_lay_mau,
            'Tiep_xuc_ca_duong_tinh':Tiep_xuc_ca_duong_tinh
              }
def multi_patient(document):
    i = 0
    documents = []
    doc = ''
    for paragraph in document.paragraphs:
        # print(paragraph.text)
        if 'Thông tin ca bệnh' in paragraph.text:
            documents.append(doc)
            doc = Document()
            print('changeeeee')
            continue
        # print(type(doc))
        if type(doc) != str:
            doc.add_paragraph(paragraph.text)
    # for d in range(1,len(documents)):
    #     print(len(documents))
    #     print(single_patient(documents[d]))
    print(documents[1])
multi_patient(document)




# with open(file_path, 'r', encoding= 'utf-8') as f:
#     for line in f:
#         # print(os.path.realpath('./'))path
#         print(line)
#         path = '/Users/user/Downloads/BÁO CÁO FILE WORD/'+line[15:-1]
#         print('here',path)
#         document = Document(path)
#         print(single_patient(document))

# document = Document(file_path)
# print(single_patient(document))
# multi_patient(document)
