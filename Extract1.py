import pandas as pd
import os
import re
from docx import Document
from collections import OrderedDict
from datetime import datetime
from datetime import timedelta


#path docx file
# file_path = '/Users/user/Downloads/BÁO CÁO FILE WORD/BC 24 QUẬN HUYỆN TỪ 1-7/QUẬN TÂN BÌNH/BN000_LÊ THỊ BÍCH TRÂM_30062021.docx'
# file_path = '/Users/user/Downloads/BÁO CÁO FILE WORD/BN0000__HƯNG YÊN_TRẦN VĂN TĂNG_030721.docx'
# file_path = '/Users/user/Downloads/BÁO CÁO FILE WORD/BC CHUỖI CHỢ BÌNH ĐIỀN - NHÓM 2+3/BN0000_ LÊ LÂM THỌ_ĐINH THỊ TRIỀU_24062021.docx'
file_path = '/Users/user/Downloads/covid_path_split_files/arr_path_1.txt'
# file_path = '/Users/user/Downloads/BÁO CÁO FILE WORD/BC 24 QUẬN HUYỆN TỪ 1-7/HUYỆN HÓC MÔN/BN0000_HM_LÊ THỊ HOÀNG_030721.docx'
# file_path = '/Users/user/Downloads/BÁO CÁO FILE WORD/ BC CHUỖI VỰA VE CHAI/BN0000_ĐẶNG NGỌC PHƯƠNG_220621_NHÓM 4.docx'
# file_path = '/Users/user/Downloads/BÁO CÁO FILE WORD/BC CHUỖI CHƯA XÁC ĐỊNH/BN00000_NGUYÊN THIÊN LỘC_26062021_BẰNG_N3.docx'
#multi
# file_path = '/Users/user/Downloads/BÁO CÁO FILE WORD/ BC CHUỖI VỰA VE CHAI/BN0000_BN0000_01_LƯƠNG THỊ THANH THUÝ_NGUYỄN THỊ LỆ THUỶ_01072021.docx'


#create document object
# document = Document(file_path)

entry_dichte = False
entry_dichte2 = False
VN_regex_cap = "ẮẰẲẴẶĂẤẦẨẪẬÂÁÀÃẢẠĐẾỀỂỄỆÊÉÈẺẼẸÍÌỈĨỊỐỒỔỖỘÔỚỜỞỠỢƠÓÒÕỎỌỨỪỬỮỰƯÚÙỦŨỤÝỲỶỸỴ"
VN_regex_norm = "áàảãạăắằẳẵặâấầẩẫậéèẻẽẹêếềểễệóòỏõọôốồổỗộơớờởỡợíìỉĩịúùủũụưứừửữự"
# date_regex = "[0-9]{1,2}/[0-9]{1,2}(?:\/[0-9]{4})?"
date_regex = "[0-9]{1,2}/[0-1]{0,1}[0-9]{0,1}(?:\/[0-9]{4})?"
date_regex_check1 = "[0-9]{1,2}/[0-1]{0,1}[0-9]{0,1}/[0-9]{4}"
date_regex_check2 = "[0-3]{0,1}[0-9]{0,1}/[0-1]{0,1}[0-9]{0,1}"
# date_regex = "[0-9]{1,2}/[0-1]{0,1}[0-9]{0,1}/[0-9]{4}"
prefix_date_regex = '(?:lấy[^.]*?'+date_regex+')|(?:[Ll]ần.*?'+date_regex+')|(?:'+date_regex+'[^\.]*?lấy mẫu)'
BN_regex = "(?:BN ?\d+)|(?:BN (?:(?:[A-Z"+VN_regex_cap+"]{1,})\s?){2,5})|(?:BN (?:(?:[A-Z"+VN_regex_cap+"][a-z"+VN_regex_norm+"]{1,})\s?){2,5})"


def extract_Ngay_duong_tinh(paragraph):
    regex = "(?:kết quả.*?dương tính[^\.]+?"+date_regex+")|(?:"+date_regex+"[^\./]+kết quả.*?dương tính)"
    regex = re.compile(regex,flags=re.I)
    list_match = None
    arr = []
    if entry_dichte:
        return list_match
    else:
        if regex.search(paragraph.text):
            list_match = regex.findall(paragraph.text)
            print('ngay_duong_tinh',list_match)
            for match in list_match:
                arr = re.compile(date_regex).findall(match)
            # list_match = list(OrderedDict.fromkeys(list_match))
            if len(arr) > 0:
                return arr[-1]
            else:
                return list_match
        else:
            regex_ngay_lay_mau = re.compile(prefix_date_regex)
            if regex_ngay_lay_mau.search(paragraph.text):
                arr = extract_Ngay_lay_mau(paragraph)
                if len(arr) > 0:
                    print('arr',arr[-1])
                    if(len(arr[-1])<= 2):
                        time = datetime.strptime(arr[-1],'%d') + timedelta(days=1)
                        return time.strftime('%d')
                    elif (len(arr[-1])<=5):
                        time = datetime.strptime(arr[-1], '%d/%m') + timedelta(days=1)
                        return time.strftime('%d/%m')
                    else:
                        time = datetime.strptime(arr[-1], '%d/%m/%Y') + timedelta(days=1)
                        return time.strftime('%d/%m/%Y')
                else:
                    return list_match
    return list_match

def extract_Dich_te(paragraph):
    regex = "[Dd]ịch [Tt]ễ:?.*"
    regex = re.compile(regex)
    global entry_dichte
    global entry_dichte2
    # print('entry',entry_dichte)
    if (regex.search(paragraph.text) != None ) or entry_dichte:
        print(entry_dichte)
        if len(paragraph.text[paragraph.text.find(':')+1:].strip()) == 0 or entry_dichte:
            print("co dau xuong dong")
            entry_dichte = True
            if re.compile('[+]').search(paragraph.text) and entry_dichte:
                entry_dichte2 = True
                if entry_dichte:
                    return paragraph.text
                else:
                    return None
            else:
                # print
                # entry_dichte = False
                if entry_dichte2:
                    entry_dichte = False
                if entry_dichte and entry_dichte2 is False:
                    if regex.search(paragraph.text) is None:
                        entry_dichte = False
                        return paragraph.text
        else:
            if(paragraph.text.find(':')):
                entry_dichte = False
                iter = paragraph.text.find(':')
                return paragraph.text[iter+1:].strip()
    return None
def extract_Ngay_lay_mau(paragraph):
    # regex = "([Dd]ương tính)"
    global entry_dichte
    regex = re.compile(prefix_date_regex)
    arr = []
    # print('entry',entry_dichte)
    if entry_dichte:
        print('Đang xét dịch tễ')
    else:
        if regex.search(paragraph.text):
            # regex = re.compile(prefix_date_regex)
            list_match = regex.findall(paragraph.text)
            print('ngay_lay_mau',list_match)
            for match in list_match:
                if re.compile(date_regex_check1).search(match):
                    arr.extend(re.compile(date_regex).findall(match))
                elif re.compile(date_regex_check2).search(match):
                    if re.compile('[Nn]gày').search(match):
                        arr.extend(re.compile(date_regex).findall(match))
            return arr
    return None
def extract_Tiep_xuc_ca_duong_tinh(paragraph):
    regex = "(?:[Dd]ương tính)|(?:[Tt]heo [Dd]iện)"
    # ([Tt]iếp xúc)
    regex = re.compile(regex)
    if regex.search(paragraph.text):
        list_match = re.compile(BN_regex).findall(paragraph.text)
        # list_match = list(OrderedDict.fromkeys(list_match))
        print('Tiep xuc',list_match)
        if len(list_match) == 0:
            return None
        else:
            return list_match
    return None

def single_patient(document):
    Ngay_lay_mau = []
    Ngay_xet_nghiem_duong_tinh = ''
    Dich_te = []
    Tiep_xuc_ca_duong_tinh = []
    i = 0
    Thong_tin_ca_benh = []
    for paragraph in document.paragraphs:
        if 'Thông tin ca bệnh' in paragraph.text:
            i += 1
        if i > 0:
            print(paragraph.text)
            Thong_tin_ca_benh.append(paragraph)
        if 'Lịch sử đi lại và tiền sử' in paragraph.text:
            i = 0
            break
    # print(Thong_tin_ca_benh)
    for paragraph in Thong_tin_ca_benh:

        res = extract_Dich_te(paragraph)
        if res != None:
            Dich_te.append(res)

        res = extract_Ngay_duong_tinh(paragraph)
        if res != None:
            Ngay_xet_nghiem_duong_tinh = res

        res = extract_Ngay_lay_mau(paragraph)
        if res != None:
            Ngay_lay_mau.extend(res)

        res = extract_Tiep_xuc_ca_duong_tinh(paragraph)
        if res != None:
            Tiep_xuc_ca_duong_tinh = res
    return {'Ngay_xet_nghiem_duong_tinh':Ngay_xet_nghiem_duong_tinh,
              'Dich_te':Dich_te,
              'Ngay_lay_mau':Ngay_lay_mau,
            'Tiep_xuc_ca_duong_tinh':Tiep_xuc_ca_duong_tinh
              }

#
# def multi_patient(document):
#     i = 0
#     documents = []
#     doc = ''
#     for paragraph in document.paragraphs:
#         # print(paragraph.text)
#         if 'Thông tin ca bệnh' in paragraph.text:
#             documents.append(doc)
#             doc = Document(docx=None)
#             print('changeeeee')
#             continue
#         # print(type(doc))
#         if 'Các hoạt động đã triển khai' in paragraph.text:
#             documents.append(doc)
#             break;
#         if type(doc) != str:
#             doc.add_paragraph(paragraph.text)
#
#     for d in range(1,len(documents)):
#         print(len(documents))
#         print(single_patient(documents[d]))
    # print(documents[1].paragraphs.text)
    # for p in documents[4].paragraphs:
    #     print(p.text)
# multi_patient(document)


# run multiple single docx
i = 1;
with open(file_path, 'r', encoding= 'utf-8') as f:
    for line in f:
        # print(os.path.realpath('./'))path
        # print(line)
        path = '/Users/user/Downloads/BÁO CÁO FILE WORD/'+line[15:-1]
        print('here',path)
        document = Document(path)

        print('\n',single_patient(document),'\n')
        print(i)
        i+=1

# test single docx
# document = Document(file_path)
# print(single_patient(document))


# regex = re.compile(BN_regex)
# list_match = regex.findall('Ngày lấy mẫu xét nghiệm: lấy mẫu lần 1 vào trưa ngày 21/06/2021 tại KCL là  trường học quận 1 (BN không biết tên KCL và địa chỉ) do tiếp xúc gần với BN12399 Lê Thị Ngọc Hương.')
# print(list_match)

# def extract_Dich_te(paragraph):
#     regex = "[Dd]ịch [Tt]ễ:?.*"
#     regex = re.compile(regex)
#     if regex.search(paragraph.text):
#         if re.compile('\n').search(paragraph.text):
#             return None
#         else:
#             if(paragraph.text.find(':')):
#                 iter = paragraph.text.find(':')
#                 return paragraph.text[iter:].strip()
#     return None

