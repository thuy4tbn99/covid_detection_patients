# -*- coding: utf-8 -*-
import pandas as pd
import os
import json
from pathlib import Path, PureWindowsPath
import re
from docx import Document
from collections import OrderedDict
from datetime import datetime
from datetime import timedelta


#path docx file
# file_path = '/Users/user/Downloads/covid_path_split_files/arr_path_1.txt'
file_path = '/Users/user/Downloads/splitted_files/normal_single.txt'

#create document object
# document = Document(file_path)

entry_dichte = False
entry_dichte2 = False
da_cach_ly = False
VN_regex_cap = "ẮẰẲẴẶĂẤẦẨẪẬÂÁÀÃẢẠĐẾỀỂỄỆÊÉÈẺẼẸÍÌỈĨỊỐỒỔỖỘÔỚỜỞỠỢƠÓÒÕỎỌỨỪỬỮỰƯÚÙỦŨỤÝỲỶỸỴ"
VN_regex_norm = "áàảãạăắằẳẵặâấầẩẫậéèẻẽẹêếềểễệóòỏõọôốồổỗộơớờởỡợíìỉĩịúùủũụưứừửữựýỳỹỷyđ"
date_regex = "[0-9]{1,2}/[0-1]{0,1}[0-9]{0,1}(?:\/[0-9]{4})?"
date_regex_check1 = "[0-9]{1,2}/[0-1]{0,1}[0-9]{0,1}/[0-9]{4}"
date_regex_check2 = "[0-3]{0,1}[0-9]{0,1}/[0-1]{0,1}[0-9]{0,1}"
prefix_date_regex = '(?:lấy[^.]*?'+date_regex+')|(?:[Ll]ần.*?'+date_regex+')|(?:'+date_regex+'[^\.)]*?lấy mẫu)|(?:[Ll][0-9].*?'+date_regex+')'
prefix_date_regex2 = ''
BN_regex = "(?:BN[ _]?\d+)|(?:BN[ _]?(?:(?:[A-Z"+VN_regex_cap+"]{1,})\s?){2,5})|(?:BN[ _]?(?:(?:[A-Z"+VN_regex_cap+"][a-z"+VN_regex_norm+"]{1,})\s?){2,5})|(?:[Bb]ệnh nhân ?(?:(?:[A-Z"+VN_regex_cap+"]{1,}?\s)){2,5})"
BN_regex2 = "(?:[Ff]0[ _]?(?:(?:[A-Z"+VN_regex_cap+"]{1,})\s?){2,5})|(?:[Ff]0[ _]?(?:(?:[A-Z"+VN_regex_cap+"][a-z"+VN_regex_norm+"]{1,})\s?){2,5})"

def docx_to_string(docx_file):
    try:
        # print(docx_file)
        document = Document(docx_file)
    except:
        print("Unable to read file")

        return ""
    return "\n".join([paragraph.text for paragraph in document.paragraphs])
def extract_sections(document_string, section):
   """"Return the string which contains the content of each section"""
   begin = ""
   end = ""
   if section == 1:
      begin = document_string.find("Thông tin ca bệnh")
      end = document_string.find("Lịch sử đi lại và tiền sử tiếp xúc và triệu chứng lâm sàng")
   elif section == 2:
      begin = document_string.find("Lịch sử đi lại và tiền sử tiếp xúc và triệu chứng lâm sàng")
      end = document_string.find("Các hoạt động đã triển khai")
   elif section == 3:
      begin = document_string.find("Các hoạt động đã triển khai")
      end = len(document_string)
   section_string = document_string[begin:end]
   return section_string


def extract_positive_date(block_text):
    regex = "(?:kết quả.*?dương tính[^\.]+?"+date_regex+")|(?:"+date_regex+"[^\./]+kết quả.*?dương tính)"
    regex = re.compile(regex,flags=re.I)
    list_match = None
    arr = []
    if entry_dichte:
        return list_match
    else:
        if regex.search(block_text):
            # print (block_text)
            list_match = regex.findall(block_text)
            print('ngay_duong_tinh',list_match)
            for match in list_match:
                arr = re.compile(date_regex).findall(match)
            # list_match = list(OrderedDict.fromkeys(list_match))
            if len(arr) > 0:
                return arr[-1]
            else:
                return list_match
        else:
            regex_ngay_lay_mau = re.compile(prefix_date_regex)
            if regex_ngay_lay_mau.search(block_text):
                # print(block_text)
                arr = extract_test_date(block_text)
                if len(arr) > 0:
                    print('arr',arr[-1])
                    if(len(arr[-1])<= 2):
                        time = datetime.strptime(arr[-1],'%d') + timedelta(days=1)
                        return time.strftime('%d')
                    else:
                        time = datetime.strptime(arr[-1], '%d/%m/%Y') + timedelta(days=1)
                        return time.strftime('%d/%m/%Y')
                else:
                    return list_match
    return list_match
def validate_test_dates(arr):
    valid_arr = []
    for d in arr:
        try:
            if(len(d)<= 2):
                time = datetime.strptime(d,'%d')
                valid_arr.append(time.strftime('%d'))
            elif (len(d)<=5):
                time = datetime.strptime(d, '%d/%m')
                valid_arr.append(time.strftime('%d/%m') + '/2021')
                # return time + datetime.timedelta(days=1)
            else:
                time = datetime.strptime(d, '%d/%m/%Y')
                valid_arr.append(time.strftime('%d/%m/%Y'))
        except:
            pass
    return valid_arr

def extract_epidemiology(block_text):
    regex = "[Dd]ịch [Tt]ễ.*:.*"
    regex = re.compile(regex)
    global entry_dichte
    global entry_dichte2
    # print('entry',entry_dichte)
    if (regex.search(block_text) != None ) or entry_dichte:
        if len(block_text[block_text.find(':')+1:].strip()) == 0 or entry_dichte:
            print("co dau xuong dong")
            entry_dichte = True
            if re.compile('[+]').search(block_text) and entry_dichte:
                entry_dichte2 = True
                if entry_dichte:
                    return block_text
                else:
                    return None
            else:
                # print
                # entry_dichte = False
                print(entry_dichte2)
                if entry_dichte2:
                    entry_dichte2 = False
                    entry_dichte = False
                if entry_dichte and entry_dichte2 is False:
                    if regex.search(block_text) is None:
                        entry_dichte = False
                        return block_text
                    else:
                        entry_dichte = True
                        return None
        else:
            if(block_text.find(':')):
                entry_dichte = False
                iter = block_text.find(':')
                return block_text[iter+1:].strip()
    return None
def extract_test_date(block_text):
    # regex = "([Dd]ương tính)"
    global entry_dichte
    regex = re.compile(prefix_date_regex)
    arr = []
    # print('entry',entry_dichte)
    if entry_dichte:
        print('Đang xét dịch tễ')
    else:
        if regex.search(block_text):
            # print (block_text)
            # regex = re.compile(prefix_date_regex)
            list_match = regex.findall(block_text)
            print('ngay_lay_mau',list_match)
            for match in list_match:
                if re.compile(date_regex_check1).search(match):
                    arr.extend(re.compile(date_regex).findall(match))
                elif re.compile(date_regex_check2).search(match):
                    if re.compile("(?:[Nn]gày)|(?:[Ll]ần ?\d{1} ?: ?"+date_regex_check2+")").search(match):
                        arr.extend(re.compile(date_regex).findall(match))
            return validate_test_dates(arr)
    return None
def extract_positive_case_contact(block_text):
    regex = "(?:[Dd]ương tính)|(?:[Tt]heo [Dd]iện)|(?:[Tt]iếp [Xx]úc (?:[Gg]ần)?)|(?:[Ll]iên quan)"
    # ([Tt]iếp xúc)
    regex = re.compile(regex)
    if regex.search(block_text):
        if re.compile("[Bb]ệnh ?[Nn]hân:").search(block_text) is None:
            list_match = re.compile(BN_regex+"|"+BN_regex2).findall(block_text)
            # list_match = list(OrderedDict.fromkeys(list_match))
            print('Tiep xuc',list_match)
            if len(list_match) == 0:
                return None
            else:
                return list_match
    return None
def extract_positive_place(block_text):
    regex_cach_ly = "(?:(?:chuyển.*)?[Cc][Áá][Cc][Hh] [Ll][Yy].*(?:do))"
    regex = "(?:[Pp]hong [Tt][oỏ][aả])|(?:[Dd]ương tính)|(?:[Tt]heo [Dd]iện)|(?:DƯƠNG TÍNH)|"+regex_cach_ly+"|(?:[Tt]iếp [Xx]úc (?:[Gg]ần)?)"
    regex = re.compile(regex)
    global da_cach_ly
    if regex.search(block_text) and da_cach_ly is False:
        print('Nguon lay',block_text)
        if re.compile("(?:[Tt]iếp [Xx]úc (?:[Gg]ần)?)|(?:[Tt]rong khu cách ly)|(?:F1)|(?:F0)|"+regex_cach_ly).search(block_text):
            print('k')
            return 'Cách ly'
        elif re.compile("[Pp]hong [Tt][oỏ][aả]").search(block_text):
            # print(re.compile("(?:[Pp]hong [Tt][oỏ][aả])").findall(block_text))
            print('e')
            if re.compile("(?:[Gg]ần) (?:(?:(?:[a-z"+VN_regex_norm+"]+) ){1,4})(?:[Pp]hong [Tt][oỏ][aả])").search(block_text) is None:
                return 'Cách ly'
            elif re.compile("trong (?:(?:(?:[a-z"+VN_regex_norm+"]+) ){1,4})(?:[Pp]hong [Tt][oỏ][aả])").search(block_text):
                return 'Cách ly'
        elif re.compile(BN_regex+"|"+BN_regex2+"|(?:xử lý theo quy trình chống dịch)").search(block_text):
            print('d')
            return 'Cách ly'
        else:
            return 'Cộng đồng'
    return None
# def extract_epidemiological_info(text_block):
#     epi_info = {'epidemiology': [], 'positve_case_contact': ''}
#     if extract_positive_date(text_block) != None:
#         # print(extract_Ngay_duong_tinh(paragraph))
#         epi_info['positive_date'] = extract_positive_date(text_block)
#     if extract_epidemiology(text_block) != None:
#         # print(extract_Dich_te(paragraph))
#         epi_info['epidemiology'].append(extract_epidemiology(text_block))
#     if extract_test_date(text_block) != None:
#         # print(extract_Ngay_lay_mau(paragraph))
#         epi_info['test_dates'] = extract_test_date(text_block)
#     if extract_positive_case_contact(text_block) != None:
#         epi_info['positve_case_contact'] = extract_positive_case_contact(text_block)
#     return epi_info
def single_patient(block_text):
    Ngay_lay_mau = []
    Ngay_xet_nghiem_duong_tinh = ''
    Dich_te = []
    Tiep_xuc_ca_duong_tinh = []
    i = 0
    Nguon_lay_nhiem = ''
    global da_cach_ly

# handle data:
    da_cach_ly = False
    if Nguon_lay_nhiem == '':
        Nguon_lay_nhiem = 'Không rõ thông tin'

    Ngay_lay_mau = list(OrderedDict.fromkeys(Ngay_lay_mau))

    if len(Ngay_lay_mau) == 0 and len(Ngay_xet_nghiem_duong_tinh) != 0:
        Ngay_lay_mau.append(Ngay_xet_nghiem_duong_tinh)

    return {'Ngay_xet_nghiem_duong_tinh':Ngay_xet_nghiem_duong_tinh,
              'Dich_te':Dich_te,
              'Ngay_lay_mau':Ngay_lay_mau,
            'Tiep_xuc_ca_duong_tinh':Tiep_xuc_ca_duong_tinh,
            'Nguon lay':Nguon_lay_nhiem
              }

def extract_epidemiological_info(text_block):
    global da_cach_ly
    Ngay_lay_mau = []
    Ngay_xet_nghiem_duong_tinh = ''
    Dich_te = []
    Tiep_xuc_ca_duong_tinh = []
    Nguon_lay_nhiem = ''
    epi_info = {'epidemiology': [], 'positve_case_contact': '', 'test_dates': [], 'positive_place': '' }

    res = extract_positive_date(text_block)
    if res != None:
        epi_info['positive_date'] = res

    res = extract_epidemiology(text_block)
    if res != None:
        epi_info['epidemiology'].append(res)

    res = extract_test_date(text_block)
    if res != None:
        epi_info['test_dates'].extend(res)

    res = extract_positive_case_contact(text_block)
    if res != None:
        epi_info['positve_case_contact'] = res

    res = extract_positive_place(text_block)
    if (res == 'Cách ly'):
        da_cach_ly = True
    if res != None:
        epi_info['positive_place'] = res


    return epi_info

