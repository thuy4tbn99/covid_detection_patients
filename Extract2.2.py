# -*- coding: utf-8 -*-
import pandas as pd
import os
import json
from pathlib import Path, PureWindowsPath
import re
from docx import Document
from collections import OrderedDict
from datetime import datetime
from datetime import timedelta


#path docx file
# file_path = '/Users/user/Downloads/splitted_files/normal_single.txt'
file_path = '/Users/user/Downloads/splitted_files/quick_report2.txt'
# file_path = '/Users/user/Downloads/baocao_covid/04-07-2021-20210707T164057Z-001/04-07-2021_review1.txt'
#multi
# file_path = '/Users/user/Downloads/BÁO CÁO FILE WORD/ BC CHUỖI VỰA VE CHAI/BN0000_BN0000_01_LƯƠNG THỊ THANH THUÝ_NGUYỄN THỊ LỆ THUỶ_01072021.docx'


#create document object
# document = Document(file_path)

entry_dichte = False
entry_dichte2 = False
da_cach_ly = False
VN_regex_cap = "ẮẰẲẴẶĂẤẦẨẪẬÂÁÀÃẢẠĐẾỀỂỄỆÊÉÈẺẼẸÍÌỈĨỊỐỒỔỖỘÔỚỜỞỠỢƠÓÒÕỎỌỨỪỬỮỰƯÚÙỦŨỤÝỲỶỸỴ"
VN_regex_norm = "áàảãạăắằẳẵặâấầẩẫậéèẻẽẹêếềểễệóòỏõọôốồổỗộơớờởỡợíìỉĩịúùủũụưứừửữựýỳỹỷyđ"
# date_regex = "[0-9]{1,2}/[0-9]{1,2}(?:\/[0-9]{4})?"
date_regex = "[0-9]{1,2}\/[0-1]{0,1}[0-9]{0,1}(?:\/[0-9]{4})?"
date_regex_check1 = "[0-9]{1,2}/[0-1]{0,1}[0-9]{0,1}/[0-9]{4}"
date_regex_check2 = "[0-3]{0,1}[0-9]{0,1}/[0-1]{0,1}[0-9]{0,1}"
# date_regex = "[0-9]{1,2}/[0-1]{0,1}[0-9]{0,1}/[0-9]{4}"
regex_positive_date = "(?:(?:kết quả)?[^\.\n0-9]*[Dd]ương tính[^\.\n(]+?" + date_regex + ")|(?:" + date_regex + "[^\.\n0-9)]+(?:kết quả)?[^\.\n)]*?[Dd]ương tính)"
prefix_date_regex = "(?:lấy[^.]*?"+date_regex+")|(?:[Ll]ần.*?"+date_regex+")|(?:"+date_regex+"[^\.)]*?lấy mẫu)|(?:[Ll][0-9].*?"+date_regex+")"
prefix_date_regex2 = "(?:\+? ?(?:"+date_regex+") ?:)|(?:(?:\+? ?(?:"+date_regex+") ?) (?:[Tt]ại))"
prefix_date_regex3 = "(?:[Xx]ét nghiệm [^\.,\n]*?"+date_regex+"[^\.,]*?)|(?:XN [^\.,]*? "+date_regex+"[^\.)]*?)|(?:"+date_regex+"[^\.,)]*?XN)|(?:"+date_regex+"[^\.,\n]*?[Xx]ét nghiệm)"
prefix_date_regex4 = "(?:"+date_regex+"[^\.,)]*?XN)"
BN_regex = "(?:BN[ _]?\d+)|(?:BN[ _]?(?:(?:[A-Z"+VN_regex_cap+"]{1,})\s?){2,5})|(?:BN[ _]?(?:(?:[A-Z"+VN_regex_cap+"][a-z"+VN_regex_norm+"]{1,})\s?){2,5})|(?:[Bb]ệnh nhân ?(?:(?:[A-Z"+VN_regex_cap+"]{1,}?\s)){2,5})"
BN_regex2 = "(?:[Ff]0[ _]?(?:(?:[A-Z"+VN_regex_cap+"]{1,})\s?){2,5})|(?:[Ff]0[ _]?(?:(?:[A-Z"+VN_regex_cap+"][a-z"+VN_regex_norm+"]{1,})\s?){2,5})"

def docx_to_string(docx_file):
    try:
        # print(docx_file)
        document = Document(docx_file)
    except:
        print("Unable to read file")
        return ""
    return "\n".join([paragraph.text for paragraph in document.paragraphs])
def extract_sections(document_string, section):
   """"Return the string which contains the content of each section"""
   begin = ""
   end = ""
   if section == 1:
      begin = document_string.find("Thông tin ca bệnh")
      end = document_string.find("Lịch sử đi lại và tiền sử tiếp xúc và triệu chứng lâm sàng")

   elif section == 2:
      begin = document_string.find("Lịch sử đi lại và tiền sử tiếp xúc và triệu chứng lâm sàng")
      end = document_string.find("Các hoạt động đã triển khai")

   elif section == 3:
      begin = document_string.find("Các hoạt động đã triển khai")
      end = len(document_string)

   section_string = document_string[begin:end]
   # return section_string
   return document_string


def extract_positive_date(block_text):
    regex = re.compile(regex_positive_date,flags=re.I)
    list_match = None
    arr = []
    if entry_dichte:
        return list_match
    else:
        if regex.search(block_text):
            list_match = regex.findall(block_text)
            print('ngay_duong_tinh',list_match)
            for match in list_match:
                if re.compile("tại ?[0-9]{1,2}\/[0-1]{0,1}[0-9]{0,1}(?:\/[0-9]{4})?").search(match) is None:
                    arr = re.compile(date_regex).findall(match)
                    arr = validate_test_dates(arr)
            if arr is not None:
                if len(arr) > 0:
                    return arr[-1]
                else:
                    return None
        else:
            regex_ngay_lay_mau = re.compile(prefix_date_regex+"|"+prefix_date_regex2+"|"+prefix_date_regex3)
            if regex_ngay_lay_mau.search(block_text):
                # print(block_text)
                arr = extract_test_date(block_text)
                if arr is not None:
                    if len(arr) > 0:
                        print('arr',arr[-1])
                        if(len(arr[-1])<= 2):
                            time = datetime.strptime(arr[-1],'%d') + timedelta(days=1)
                            return time.strftime('%d')
                        else:
                            time = datetime.strptime(arr[-1], '%d/%m/%Y') + timedelta(days=1)
                            return time.strftime('%d/%m/%Y')
                    else:
                        return list_match
    return list_match

def extract_epidemiology(block_text):
    regex = "[Dd]ịch [Tt]ễ.*:.*"
    regex = re.compile(regex)
    global entry_dichte
    global entry_dichte2
    # print('entry',entry_dichte)
    if (regex.search(block_text) != None ) or entry_dichte:
        if len(block_text[block_text.find(':')+1:].strip()) == 0 or entry_dichte:
            print("co dau xuong dong")
            entry_dichte = True
            if re.compile('[+]').search(block_text) and entry_dichte:
                entry_dichte2 = True
                if entry_dichte:
                    return block_text
                else:
                    return None
            else:
                print(entry_dichte2)
                if entry_dichte2:
                    entry_dichte2 = False
                    entry_dichte = False
                if entry_dichte and entry_dichte2 is False:
                    if regex.search(block_text) is None:
                        entry_dichte = False
                    else:
                        entry_dichte = True
                        return None
        else:
            if(block_text.find(':')):
                entry_dichte = False
                iter = block_text.rfind(':')
                return block_text[iter+1:].strip()
    return None

def extract_positive_case_contact(block_text):
    regex = "(?:[Dd]ương tính)|(?:[Tt]heo [Dd]iện)|(?:[Tt]iếp [Xx]úc (?:[Gg]ần)?)|(?:[Ll]iên quan)"
    # ([Tt]iếp xúc)
    regex = re.compile(regex)
    if regex.search(block_text):
        if re.compile("[Bb]ệnh ?[Nn]hân:").search(block_text) is None:
            list_match = re.compile(BN_regex+"|"+BN_regex2).findall(block_text)
            print('Tiep xuc',list_match)
            for match in list_match:
                if re.compile("(?:TIẾP XÚC)|(?:[Tt]iếp [Xx]úc)").search(match):
                    list_match.remove(match)
            if len(list_match) == 0:
                return None
            else:
                return list_match
    return None

def extract_positive_place(block_text):
    regex_cach_ly = "(?:(?:chuyển.*)?[Cc][Áá][Cc][Hh] [Ll][Yy].*(?:do))"
    regex = "(?:[Pp]hong [Tt][oỏ][aả])|(?:[Dd]ương tính)|(?:[Tt]heo [Dd]iện)|(?:DƯƠNG TÍNH)|"+regex_cach_ly+"|(?:[Tt]iếp [Xx]úc (?:[Gg]ần)?)"
    regex = re.compile(regex)
    global da_cach_ly
    if regex.search(block_text) and da_cach_ly is False:
        print('Nguon lay',block_text)
        if re.compile("(?:[Tt]iếp [Xx]úc (?:[Gg]ần)?)|(?:[Tt]rong ?(?:đến)? khu cách ly)|(?:F1)|(?:F0)|"+regex_cach_ly).search(block_text):
            print('o')
            if re.compile("không rõ F0").findall(block_text):
                return 'Cộng đồng'
            else:
                return 'Cách ly'
        elif re.compile("[Pp]hong [Tt][oỏ][aả]").search(block_text):
            print('k')
            if re.compile("(?:[Gg]ần) (?:(?:(?:[a-z"+VN_regex_norm+"]+) ){1,4})(?:[Pp]hong [Tt][oỏ][aả])").search(block_text) is None:
                return 'Cách ly'
            elif re.compile("trong (?:(?:(?:[a-z"+VN_regex_norm+"]+) ){1,4})(?:[Pp]hong [Tt][oỏ][aả])").search(block_text):
                return 'Cách ly'
        elif re.compile(BN_regex+"|"+BN_regex2+"|(?:xử lý theo quy trình chống dịch)").search(block_text):
            print('e')
            return 'Cách ly'
        else:
            return 'Cộng đồng'
    return None

def extract_test_date(block_text):
    global entry_dichte
    regex = re.compile(prefix_date_regex)
    arr = []
    # print('entry',entry_dichte)
    if entry_dichte:
        print('Đang xét dịch tễ')
    else:
        if regex.search(block_text):
            list_match = regex.findall(block_text)
            print('ngay_lay_mau',list_match)
            for match in list_match:
                if re.compile(date_regex_check1).search(match):
                    arr.extend(re.compile(date_regex).findall(match))
                elif re.compile(date_regex_check2).search(match):
                    if re.compile("(?:[Nn]gày)|(?:[Ll]ần ?\d{1} ?: ?"+date_regex_check2+")").search(match):
                        arr.extend(re.compile(date_regex).findall(match))
            return validate_test_dates(arr)
        # elif re.compile("(?:\+? ?(?:"+date_regex+") ?:)").search(block_text):
        #     print('ngay :')
        #     print(block_text)
        #     list_match = re.compile("(?:\+? ?(?:"+date_regex+") ?:)").findall(block_text)
        #     for match in list_match:
        #         arr.extend(re.compile(date_regex).findall(match))
        #     return validate_test_dates(arr)
        elif re.compile("(?:\+? ?(?:"+date_regex+") ?) (?:[Tt]ại)").search(block_text):
            print('ngay : tai')
            list_match = re.compile("(?:\+? ?(?:"+date_regex+") ?) (?:[Tt]ại)").findall(block_text)
            for match in list_match:
                arr.extend(re.compile(date_regex).findall(match))
            return validate_test_dates(arr)
        elif re.compile(prefix_date_regex3).search(block_text):
            print('xn')
            list_match = re.compile(prefix_date_regex3).findall(block_text)
            for match in list_match:
                arr.extend(re.compile(date_regex).findall(match))
            return validate_test_dates(arr)
    return None

def validate_test_dates(arr):
    valid_arr = []
    for d in arr:
        try:
            if(len(d)<= 2):
                time = datetime.strptime(d,'%d')
                valid_arr.append(time.strftime('%d'))
            elif (len(d)<=5):
                time = datetime.strptime(d, '%d/%m')
                valid_arr.append(time.strftime('%d/%m') + '/2021')
            else:
                time = datetime.strptime(d, '%d/%m/%Y')
                valid_arr.append(time.strftime('%d/%m/%Y'))
        except:
            pass
    return valid_arr
def extract_hospital(text_block):
    list_match = []
    regex_hos = "(?:[Cc]huyển (?:[a-z"+VN_regex_norm+"]* ?){1,4}(?:bệnh viện)?(?:bv)?(?:[0-9a-z"+VN_regex_norm+"]* ?){1,})"
    regex_hos2 = "(?:cách ly tại *:? *(?:(?:(?:[0-9a-z"+VN_regex_norm+"]+) *){0,5}))"
    regex_hos3 = "(?:[Đđ]ưa *(?:[a-z"+VN_regex_norm+"]* ?){1,4}(?:bệnh viện)?(?:bv)?(?:[0-9a-z"+VN_regex_norm+"]* ?){1,})"
    hospital = ['Bệnh viện Dã chiến Củ Chi', 'Bệnh viện Bệnh Nhiệt đới TPHCM',
       'Bệnh viện điều trị COVID-19 Cần Giờ', 'TTYT Bến Cầu',
       'Bệnh viện Nhi Đồng 2 - TP HCM',
       'Bệnh viện Dã chiến Thu Dung số 1 (KTX TTGDQPAN Đại học Quốc Gia TPHCM)',
        'Bệnh viện Củ Chi', 'Bệnh viện Nhi Đồng 1',
       'Bệnh viện Huyện Củ Chi',
       'bệnh Nhiệt đới thành phố Hồ Chí Minh.', 'COVID19 Cần Giờ',
       'Bệnh viện Phạm Ngọc Thạch', 'Bệnh viện Chợ Rẫy',
        'Bệnh viện Trưng Vương',
       'Bệnh viện Nhi Đồng 2 - TP HCM ', 'Bệnh viện Trưng Vương',
       'Bệnh viện Đa khoa Thủ Đức', 'Bệnh viện dã chiến Bình Chánh',
       'Bệnh viện dã chiến quận 9', 'Bệnh viện Nguyễn Trãi',
       'Bệnh viện Đa khoa tỉnh Bình Dương', 'Bệnh viện Việt Thắng',
       'Bệnh viện dã chiến Củ Chi', 'Bệnh viện Bình Dân',
       'Bệnh viện quận 8', 'Bệnh viện quận 10',
       'Bệnh viện Dã chiến Thu Dung số 2',
       'Dã chiến Thu Dung 3', 'Bệnh viện Thủ Đức', 'ktx đhqg', 'Bệnh viện covid cần giờ' , 'Bệnh viện Thu Dung' , 'Bệnh viện Bình Chánh', 'Khu A số 1 Lê Quý Đôn, phường Đông Hòa, TP.Dĩ An, Bình Dương',
        'KTX Đại học Quốc gia']
    if re.compile(regex_hos+"|"+regex_hos2+"|"+regex_hos3).search(text_block.lower()):
        list_match.extend(re.compile(regex_hos+"|"+regex_hos2+"|"+regex_hos3).findall(text_block.lower()))
    # print('hos',list_match)
    for index, h in enumerate(hospital):
        for match in list_match:
            hos = h.strip().replace('Bệnh viện ','').lower()
            if hos in match:
                print(hos)
                return hospital[index]
            elif 'điều trị' in match:
                if re.compile("(?:BV)|(?:[Bb]ệnh viện)").search(match):
                    return 'Bệnh viện điều trị'
    return None

def extract_symptom(text_block):
    if re.compile("[Tt]riệu chứng *[^\n]*:").search(text_block):
        print(text_block)
        iter = text_block.find(":")
        return text_block[iter+1:]
    return None

def clean_info(epi_info):
    epi_info['test_dates'] = list(OrderedDict.fromkeys(epi_info['test_dates']))

    if len(epi_info['test_dates']) == 0 and len(epi_info['positive_date']) != 0:
        epi_info['test_dates'].append(epi_info['positive_date'])
    if len(epi_info['test_dates']) == 0:
        epi_info['test_dates'] = ''

    print(epi_info['positive_date'])
    if epi_info['test_dates'] != '' and epi_info['positive_date'] != '':
        datetime1 = datetime.strptime(epi_info['test_dates'][-1], '%d/%m/%Y')
        datetime2 = datetime.strptime(epi_info['positive_date'] , '%d/%m/%Y')
        if datetime1 > datetime2:
            epi_info['positive_date'] = epi_info['test_dates'][-1]
    return epi_info

def extract_epidemiological_info(text_block):
    global da_cach_ly

    epi_info = {'positve_case_contact': '', 'test_dates': [], 'positive_place': 'Không rõ thông tin', 'positive_date':'', 'hospital':'' ,'epidemiology': [],'symptom':''}

    print(extract_sections(text_block, 1))
    Thong_tin_ca_benh = extract_sections(text_block, 1).split('\n')
    # Thong_tin_ca_benh = Thong_tin_ca_benh.split('\n')
    # print(extract_sections(text_block,3).lower())
    Thong_tin_di_lai = extract_sections(text_block,2).split('\n')
    Hoat_dong_trien_khai = extract_sections(text_block,3).lower().split('\n')

    for text in Thong_tin_ca_benh:

        res = extract_epidemiology(text)
        if res != None:
            epi_info['epidemiology'].append(res)

        res = extract_positive_date(text)
        if res != None:
            epi_info['positive_date'] = res

        res = extract_test_date(text)
        if res != None:
            epi_info['test_dates'].extend(res)

        res = extract_positive_case_contact(text)
        if res != None:
            epi_info['positve_case_contact'] = res

        res = extract_positive_place(text)
        if (res == 'Cách ly'):
            da_cach_ly = True
        if res != None:
            epi_info['positive_place'] = res

        res = extract_hospital(text)
        if res != None:
            epi_info['hospital'] = res

        res = extract_symptom(text)
        if res != None:
            epi_info['symptom'] = res

    for text in Thong_tin_di_lai:
        res = extract_symptom(text)
        if res != None:
            epi_info['symptom'] = res

        res = extract_hospital(text)
        if res != None:
            epi_info['hospital'] = res

    for text in Hoat_dong_trien_khai:
        res = extract_hospital(text)
        if res != None:
            epi_info['hospital'] = res


    da_cach_ly = False

    epi_info = clean_info(epi_info)
    return epi_info

# # run multiple single docx
i = 1;
output = []
with open(file_path, 'r', encoding= 'utf-8') as f:
    for line in f:
        path = PureWindowsPath(line)
        path = Path(path)
        path1 = Path("/Users/user/Downloads")
        path2 = Path.joinpath(path1,path)
        path3 = str(path2)
        print('here',path3)
        document_string = docx_to_string(path3[:-1])
        # print('Here',line)
        # print(document_string)
        # document_string = docx_to_string(line[:-1])
        # print(extract_sections(document_string,3))
        res = extract_epidemiological_info(document_string)
        print('\n',res,'\n')
        res['Link'] = path3[:-1]
        output.append(res)
        print(i)
        i+=1
        if i == (80 + 2):
            break


df = pd.DataFrame.from_records(output)
# df.to_excel("/Users/user/Desktop/test/extract9.xlsx")
# test single docx
# file_path = "/Users/user/Downloads/baocao_covid/04-07-2021-20210707T164057Z-001/04-07-2021/BN19095_Tống Hữu Lộc.docx"
# document = Document(file_path)
# a = docx_to_string(file_path)
# print('\n',single_patient(a),'\n')




