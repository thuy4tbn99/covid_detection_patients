from docx import Document
import os
import re
import pandas as pd
import json
import bao_cao_nhanh
import argparse

from collections import OrderedDict
from datetime import datetime
from dateutil import parser
from datetime import date
from categorizer import DocumentClassifier

from pathlib import Path, PureWindowsPath
import re
from collections import OrderedDict
from datetime import datetime
from datetime import timedelta

parser = argparse.ArgumentParser(description='WILDCAT Training')
parser.add_argument('--dir', default=None,
                    type=str, metavar='DIR', help='Path to the data directory')
# parser.add_argument('-o', "--output_file", default=None,
#                     type=str, metavar='OF', help='Path to the output file')

today_date = date.today()

#-----------------------------------------------------------------------------
#1. PERSONAL INFO
def extract_single_patient(file_path):
     # get personal infor
    document = Document(file_path)
    
    lines = [paragraph.text.strip() for paragraph in document.paragraphs]
    raw_text = '\n'.join(lines)
    
    if raw_text.lower().find('cục y tế dự phòng;') == -1:
        return {'doc_type': 'baocaonhanh'}
    
    basic_info_section = extract_sections(raw_text, section=1)
    
    basic_info = extract_basic_info(basic_info_section)
    location_info = extract_location_info(basic_info_section)
    epi_info = extract_epidemiological_info(basic_info_section)

    patient_info = basic_info.copy()
    patient_info.update(location_info)
    patient_info.update(epi_info)
    
    return patient_info

def extract_sections(raw_text, section):
   """"Return the string which contains the content of each section"""
   if section == 1:
      begin = raw_text.find("Thông tin ca bệnh")
      if begin == -1:
          begin = raw_text.find("Nhận thông tin") 
      end = raw_text.find("Lịch sử đi lại và tiền sử tiếp xúc và triệu chứng lâm sàng")
   elif section == 2:
      begin = raw_text.find("Lịch sử đi lại và tiền sử tiếp xúc và triệu chứng lâm sàng")
      end = raw_text.find("Các hoạt động đã triển khai")
   elif section == 3:
      begin = raw_text.find("Các hoạt động đã triển khai")
      end = len(raw_text)
   section_string = raw_text[begin:end]
   return section_string

def extract_basic_info(text_block):
    regex_dict = {
        "hoTen" : "(Bệnh nhân:?.?([\w\sắằẳẵặăấầẩẫậâáàãảạđếềểễệêéèẻẽẹíìỉĩịốồổỗộôớờởỡợơóòõỏọứừửữựưúùủũụýỳỷỹỵẮẰẲẴẶĂẤẦẨẪẬÂÁÀÃẢẠĐẾỀỂỄỆÊÉÈẺẼẸÍÌỈĨỊỐỒỔỖỘÔỚỜỞỠỢƠÓÒÕỎỌỨỪỬỮỰƯÚÙỦŨỤÝỲỶỸỴ-]){2,}([\t\n(,]){1})",
        "maBN":"(BN\s?[0-9]+)",
        "namSinh":"(((sinh năm)|(SN)|(Sinh ngày))[:]*.?[\s\d\/]*\d{4})",
        "gioiTinh":"(\s?(nam|nữ|NAM|NỮ|Nam|Nữ))",
        "CMND":"(((nhân dân)|(CCCD)):\s?\d{8,})",
        "quocTich": "(tịch: [a-zắằẳẵặăấầẩẫậâáàãảạđếềểễệêéèẻẽẹíìỉĩịốồổỗộôớờởỡợơóòõỏọứừửữựưúùủũụýỳỷỹỵA-ZẮẰẲẴẶĂẤẦẨẪẬÂÁÀÃẢẠĐẾỀỂỄỆÊÉÈẺẼẸÍÌỈĨỊỐỒỔỖỘÔỚỜỞỠỢƠÓÒÕỎỌỨỪỬỮỰƯÚÙỦŨỤÝỲỶỸỴ\s]+)[.,\n]+",
        "SDT":"(\d{4}.?\d{3}.?\d{3})"
    }

    regex_info = {}
    for regex in regex_dict:
        if re.search(re.compile(regex_dict[regex]), text_block):
            regex_info[regex] = [m for m in re.findall(re.compile(regex_dict[regex]), text_block)][0]
            # print('rtn_dict',regex, rtn_dict[regex])
            
    if 'hoTen' not in regex_info:
        start_idx = text_block.find('Bệnh nhân:')
        end_idx = text_block.index('(', start_idx)
        regex_info['hoTen'] = [text_block[start_idx : end_idx].strip()]

    return standardalize_basic_info(regex_info)

def standardalize_basic_info(raw_dict):
    basic_info = {}
    basic_info['doc_type'] = 'bccabenh'
    basic_info["name"] = raw_dict["hoTen"][0][10:].replace('(', '').strip()

    if 'maBN' in raw_dict:
        basic_info["patient_code"] = raw_dict["maBN"].strip()
    else:
        basic_info["patient_code"] = ''

    if 'namSinh' in raw_dict:
        basic_info["birthyear"] = raw_dict["namSinh"][0][-4:].strip()
    else:
        basic_info["birthyear"] = ''

    basic_info["gender"] = raw_dict["gioiTinh"][1:][0].lower().strip()

    if "CMND" in raw_dict and len(raw_dict["CMND"][0]) > 8:
        basic_info["CMND"] = re.findall('[0-9]+', raw_dict["CMND"][0])[0].strip()
    else: 
        basic_info["CMND"] =""
 
    if 'quocTich' in raw_dict:
        start_idx = raw_dict["quocTich"].find(":")
        basic_info["nationality"] = raw_dict["quocTich"][start_idx + 1:].strip()
    else:
        basic_info["nationality"] = ''
    
    if 'SDT' in raw_dict:
        basic_info["phone"] = raw_dict["SDT"].strip()
    else:
        basic_info["phone"] = ''

    return basic_info

#-----------------------------------------------------------------------------
# 2. Extract location info

def extract_location_info(text_block):
    """"Return a json file which contains patient's information extracted from a word document"""
    #get patient's job's description
    patient_job, patient_job_location = find_job(text_block)
    #get patient's address location
    patient_address = find_address(text_block).lower()
    street, village, district, provine = split_address_normal(patient_address)
    output = {
        "job": patient_job,
        'job_loc': patient_job_location,
        "adress": patient_address,
        "street": street,
        "ward": village,
        "district": district,
        "provine": provine
    }
    return output

def find_job(section_string):
    """Return patient's job description from patient's info string"""
    pattern = re.compile(r'(?<=Nghề nghiệp:)[^\n]*(?=\n)|(?<=Tên và địa chỉ nơi làm việc:)[^\n]*(?=\n)|(?<=Tên và địa chỉ làm việc)[^\n]*(?=\n)|(?<=Địa chỉ nơi ở và nơi làm)[^\n]*(?=\n)')
    job_search = re.search(pattern, section_string)
    if job_search:
        job_info = job_search.group().strip()
        return find_job_location(job_info)
        
    return 'NA', 'NA'

def find_job_location(job_info):
    pattern = re.compile(r'(?<=tại)[^\n]+|(?<=ông ty)[^\n]+|(?<=trường)[^\n]+')
    
    if re.search(re.compile(pattern),job_info):
        job_loc = re.findall(pattern, job_info)[0].strip()
        if "ông ty" in job_info and "ông ty" not in job_loc:
            job_loc = "công ty " + job_loc
        
        if "trường" in job_info and "trường" not in job_loc:
            job_loc = "trường " + job_loc
        job_name = job_info[:-len(job_loc)].strip()
        job_name = job_name.replace('tại', '')
        return job_name, job_loc 
    
    return job_info, 'NA'
    
def find_address(document_string):
    """Return patient's address from patient's info stirng"""
    address = ""
    pattern = re.compile(r'(?<=Địa chỉ:)[^\n]+|(?<=Địa chỉ nơi ở:)[^\n]+|(?<=Địa chỉ tạm trú:)[^\n]+|(?<=Địa chỉ nơi ở hiện nay:)[^\n]+|(?<=Địa chỉ nơi)[^\n]+|(?<=Địa chỉ nơi ở và nơi làm)[^\n]+|(?<=Địa chỉ nhà:)[^\n]+')
    if re.search(re.compile(pattern),document_string):
        address = re.findall(pattern, document_string)[0]
    
    if ':' in address:
        address = address[address.find(':') + 1:].strip()    
    
    return address

def split_address_normal(address_string):
        address_string = address_string.replace('.', '')
        districts = [ 'quận 11','q11', 'quận 12','q12' ,'quận 10','q10','quận 9','q9', 'quận 4','q4' ,   'quận 6','q6',
            'quận 2','q2','quận 5', 'q5', 'quận 7', 'q7','q3' ,'quận 3',  'quận 1', 'q1' ,
            'cần giờ','củ chi','gò vấp', 'phú nhuận',  'bình thạnh', 'quận 8','q8', 'tân bình', 'nhà bè',  
            'hóc môn', 'bình chánh','thủ đức', 'tân phú','bình tân']
        villages = ['phường 16', 'p16', 'tăng nhơn phú b', 'thủ thiêm', 'linh đông', 'bình trưng tây', 'hoà thạnh', 
            'phường 07', 'p7', 'phường 08','p8', 'thảo điền', 'tân chánh hiệp', 'hiệp thành', 'thạnh an', 'vĩnh lộc a', 
            'bình thuận', 'phú xuân', 'tân an hội', 'nhơn đức', 'phường 11','p11' ,'long thới', 'hóc môn', 
            'an nhơn tây', 'phước long b', 'trường thọ', 'bình trưng đông', 'an phú đông', 'bình chiểu', 'bình mỹ',
            'cầu kho', 'thái mỹ', 'phạm văn hai', 'hưng long', 'tân hiệp', 'thạnh xuân', 'trung mỹ tây', 'hiệp bình chánh', 
            'phú thuận', 'tân hưng', 'an lạc', 'tân thới hoà', 'thạnh lộc', 'long thạnh mỹ', 'bình hưng hoà b', 'bà điểm', 
            'tân quý tây', 'bình khánh', 'phường 28', 'p28', 'hiệp phước', 'phú trung', 'bình hưng hòa', 
            'đông hưng thuận', 'phường 7','p7', 'nhị bình', 'tân kiểng', 'an lợi đông', 'linh chiểu', 'an khánh', 
            'phú hòa đông', 'phạm ngũ lão', 'bình an', 'tân thạnh tây', 'phước long a', 'bình chánh', 'linh trung', 'bình trị đông', 'bến nghé', 'long bình', 'tân thuận tây', 'tăng nhơn phú a', 'phước thạnh', 
            'tam phú', 'xuân thới thượng', 'phường 12','p12', 'trung lập thượng', 'tân sơn nhì', 'thới tam thôn', 
            'tân tạo a', 'bình trị đông a', 'phường 24','p24', 'phước bình', 'bình trị đông b', 'phước lộc', 
            'tân qúy', 'đông thạnh', 'lê minh xuân', 'tân nhựt', 'tam bình', 'phú hữu', 'phước kiển', 'tây thạnh', 
            'an thới đông', 'phường 10', 'p10', 'đa kao', 'nhà bè', 'hiệp tân', 'phạm văn cội', 'phú thọ hoà', 
            'quy đức', 'an phú tây', 'tân xuân', 'phú mỹ hưng', 'phường 27', 'p27', 'linh xuân', 'phường 26', 'p26', 
            'nguyễn thái bình', 'phước vĩnh an', 'linh tây', 'tân phong', 'bến thành', 'hiệp phú', 
            'bình hưng hoà a', 'phường 21', 'long phước', 'trung lập hạ', 'phường 17', 'p17', 'tân thạnh đông', 
            'cô giang', 'xuân thới đông', 'bình lợi', 'an phú', 'đa phước', 'phú mỹ', 'tân hưng thuận', 'bình thọ', 
            'phường 18','p18', 'tân tạo', 'phước hiệp', 'cần thạnh', 'phường 13', 'p13','tân quy', 'hiệp bình phước', 
            'phường 15', 'phường 05', 'long hòa', 'hòa phú', 'cầu ông lãnh', 'cát lái', 'thới an', 'tân thuận đông', 
            'tân thới hiệp', 'xuân thới sơn', 'thạnh mỹ lợi', 'tân kiên', 'lý nhơn', 'an lạc a', 'trung an', 'phong phú', 
            'bình hưng', 'nguyễn cư trinh', 'phường 3', 'p3', 'phường 6','p6' ,'phường 19','p19','tân túc', 'phú thạnh', 'phường 14','p14', 'tân phú trung', 'p. tân phú', 'tân phú', 'tân thới nhì', 'phường 22', 'p22',
            'tân thông hội', 'sơn kỳ', 'trung chánh', 'tân định', 'tân thới nhất', 'tam thôn hiệp', 'vĩnh lộc b', 
            'phường 25','p25', 'phường 9','p9', 'củ chi', 'tân thành', 'nhuận đức', 'long trường', 'trường thạnh', 
            'phường 1', 'p1', 'phường 2', 'p2', 'phường 4','p4']
        hcm = {'tp hcm', 'tphcm', 'tp hồ chí minh', 'thành phố hồ chí minh'}
        street = ""
        village = ""
        district = ""
        provine = ""
        for i in hcm:
                if i in address_string:
                    provine = "TP HCM"
                    address_string = address_string[:address_string.find(i)]
                    break
        for d in districts:
                if d in address_string:
                    if d == "tân phú":
                        break
                    district = d
                    provine = "TP HCM"
                    address_string = address_string[:address_string.find(d)]
                    break
        for v in villages:
            if v in address_string:
                if v == "tân phú":
                    break
                village = v
                provine = 'TP HCM'
                street= address_string[:address_string.find(v)]
                street = street.replace("xã", "")
                street = street.replace("phường", "")
                street = street.replace(",", "")
                break
        return street, village, district, provine

#------------------------------------------------------------
# Extract epidemiological info
entry_dichte = False
entry_dichte2 = False
da_cach_ly = False
VN_regex_cap = "ẮẰẲẴẶĂẤẦẨẪẬÂÁÀÃẢẠĐẾỀỂỄỆÊÉÈẺẼẸÍÌỈĨỊỐỒỔỖỘÔỚỜỞỠỢƠÓÒÕỎỌỨỪỬỮỰƯÚÙỦŨỤÝỲỶỸỴ"
VN_regex_norm = "áàảãạăắằẳẵặâấầẩẫậéèẻẽẹêếềểễệóòỏõọôốồổỗộơớờởỡợíìỉĩịúùủũụưứừửữựýỳỹỷyđ"
date_regex = "[0-9]{1,2}/[0-1]{0,1}[0-9]{0,1}(?:\/[0-9]{4})?"
date_regex_check1 = "[0-9]{1,2}/[0-1]{0,1}[0-9]{0,1}/[0-9]{4}"
date_regex_check2 = "[0-3]{0,1}[0-9]{0,1}/[0-1]{0,1}[0-9]{0,1}"
prefix_date_regex = '(?:lấy[^.]*?'+date_regex+')|(?:[Ll]ần.*?'+date_regex+')|(?:'+date_regex+'[^\.)]*?lấy mẫu)|(?:[Ll][0-9].*?'+date_regex+')'
prefix_date_regex2 = ''
BN_regex = "(?:BN[ _]?\d+)|(?:BN[ _]?(?:(?:[A-Z"+VN_regex_cap+"]{1,})\s?){2,5})|(?:BN[ _]?(?:(?:[A-Z"+VN_regex_cap+"][a-z"+VN_regex_norm+"]{1,})\s?){2,5})|(?:[Bb]ệnh nhân ?(?:(?:[A-Z"+VN_regex_cap+"]{1,}?\s)){2,5})"
BN_regex2 = "(?:[Ff]0[ _]?(?:(?:[A-Z"+VN_regex_cap+"]{1,})\s?){2,5})|(?:[Ff]0[ _]?(?:(?:[A-Z"+VN_regex_cap+"][a-z"+VN_regex_norm+"]{1,})\s?){2,5})"

def extract_epidemiological_info(text_block):
    global da_cach_ly

    epi_info = {'epidemiology': [], 'positve_case_contact': '', 'test_dates': [], 'positive_place': 'Không rõ thông tin' }

    Thong_tin_ca_benh = text_block.split('\n')

    for text in Thong_tin_ca_benh:
        res = extract_positive_date(text)
        if res != None:
            epi_info['positive_date'] = res

        res = extract_epidemiology(text)
        if res != None:
            epi_info['epidemiology'].append(res)

        res = extract_test_date(text)
        if res != None:
            epi_info['test_dates'].extend(res)

        res = extract_positive_case_contact(text)
        if res != None:
            epi_info['positve_case_contact'] = res

        res = extract_positive_place(text)
        if (res == 'Cách ly'):
            da_cach_ly = True
        if res != None:
            epi_info['positive_place'] = res

    da_cach_ly = False

    epi_info['test_dates'] = list(OrderedDict.fromkeys(epi_info['test_dates'] ))

    if len(epi_info['test_dates']) == 0 and len(epi_info['positive_date']) != 0:
        epi_info['test_dates'] .append(epi_info['positive_date'])

    return epi_info

def extract_positive_date(block_text):
    regex = "(?:kết quả.*?dương tính[^\.]+?"+date_regex+")|(?:"+date_regex+"[^\./]+kết quả.*?dương tính)"
    regex = re.compile(regex,flags=re.I)
    list_match = None
    arr = []
    if entry_dichte:
        return list_match
    else:
        if regex.search(block_text):
            # print (block_text)
            list_match = regex.findall(block_text)
            print('ngay_duong_tinh',list_match)
            for match in list_match:
                arr = re.compile(date_regex).findall(match)
            # list_match = list(OrderedDict.fromkeys(list_match))
            if len(arr) > 0:
                return arr[-1]
            else:
                return list_match
        else:
            regex_ngay_lay_mau = re.compile(prefix_date_regex)
            if regex_ngay_lay_mau.search(block_text):
                # print(block_text)
                arr = extract_test_date(block_text)
                if len(arr) > 0:
                    print('arr',arr[-1])
                    if(len(arr[-1])<= 2):
                        time = datetime.strptime(arr[-1],'%d') + timedelta(days=1)
                        return time.strftime('%d')
                    else:
                        time = datetime.strptime(arr[-1], '%d/%m/%Y') + timedelta(days=1)
                        return time.strftime('%d/%m/%Y')
                else:
                    return list_match
    return list_match

def extract_epidemiology(block_text):
    regex = "[Dd]ịch [Tt]ễ.*:.*"
    regex = re.compile(regex)
    global entry_dichte
    global entry_dichte2
    # print('entry',entry_dichte)
    if (regex.search(block_text) != None ) or entry_dichte:
        if len(block_text[block_text.find(':')+1:].strip()) == 0 or entry_dichte:
            print("co dau xuong dong")
            entry_dichte = True
            if re.compile('[+]').search(block_text) and entry_dichte:
                entry_dichte2 = True
                if entry_dichte:
                    return block_text
                else:
                    return None
            else:
                # print
                # entry_dichte = False
                print(entry_dichte2)
                if entry_dichte2:
                    entry_dichte2 = False
                    entry_dichte = False
                if entry_dichte and entry_dichte2 is False:
                    if regex.search(block_text) is None:
                        entry_dichte = False
                        return block_text
                    else:
                        entry_dichte = True
                        return None
        else:
            if(block_text.find(':')):
                entry_dichte = False
                iter = block_text.find(':')
                return block_text[iter+1:].strip()
    return None
<<<<<<< HEAD
=======

>>>>>>> 3a3ea677d45677e2a850c55fa547dfce3f852eba
def extract_positive_case_contact(block_text):
    regex = "(?:[Dd]ương tính)|(?:[Tt]heo [Dd]iện)|(?:[Tt]iếp [Xx]úc (?:[Gg]ần)?)|(?:[Ll]iên quan)"
    # ([Tt]iếp xúc)
    regex = re.compile(regex)
    if regex.search(block_text):
        if re.compile("[Bb]ệnh ?[Nn]hân:").search(block_text) is None:
            list_match = re.compile(BN_regex+"|"+BN_regex2).findall(block_text)
            # list_match = list(OrderedDict.fromkeys(list_match))
            print('Tiep xuc',list_match)
            if len(list_match) == 0:
                return None
            else:
                return list_match
<<<<<<< HEAD
    return None

def extract_test_date(block_text):
    # regex = "([Dd]ương tính)"
=======
    return None

def extract_positive_place(block_text):
    regex_cach_ly = "(?:(?:chuyển.*)?[Cc][Áá][Cc][Hh] [Ll][Yy].*(?:do))"
    regex = "(?:[Pp]hong [Tt][oỏ][aả])|(?:[Dd]ương tính)|(?:[Tt]heo [Dd]iện)|(?:DƯƠNG TÍNH)|"+regex_cach_ly+"|(?:[Tt]iếp [Xx]úc (?:[Gg]ần)?)"
    regex = re.compile(regex)
    global da_cach_ly
    if regex.search(block_text) and da_cach_ly is False:
        print('Nguon lay',block_text)
        if re.compile("(?:[Tt]iếp [Xx]úc (?:[Gg]ần)?)|(?:[Tt]rong khu cách ly)|(?:F1)|(?:F0)|"+regex_cach_ly).search(block_text):
            print('k')
            return 'Cách ly'
        elif re.compile("[Pp]hong [Tt][oỏ][aả]").search(block_text):
            # print(re.compile("(?:[Pp]hong [Tt][oỏ][aả])").findall(block_text))
            print('e')
            if re.compile("(?:[Gg]ần) (?:(?:(?:[a-z"+VN_regex_norm+"]+) ){1,4})(?:[Pp]hong [Tt][oỏ][aả])").search(block_text) is None:
                return 'Cách ly'
            elif re.compile("trong (?:(?:(?:[a-z"+VN_regex_norm+"]+) ){1,4})(?:[Pp]hong [Tt][oỏ][aả])").search(block_text):
                return 'Cách ly'
        elif re.compile(BN_regex+"|"+BN_regex2+"|(?:xử lý theo quy trình chống dịch)").search(block_text):
            print('d')
            return 'Cách ly'
        else:
            return 'Cộng đồng'
    return None

def extract_test_date(block_text):
>>>>>>> 3a3ea677d45677e2a850c55fa547dfce3f852eba
    global entry_dichte
    regex = re.compile(prefix_date_regex)
    arr = []
    # print('entry',entry_dichte)
    if entry_dichte:
        print('Đang xét dịch tễ')
    else:
        if regex.search(block_text):
            # print (block_text)
            # regex = re.compile(prefix_date_regex)
            list_match = regex.findall(block_text)
            print('ngay_lay_mau',list_match)
            for match in list_match:
                if re.compile(date_regex_check1).search(match):
                    arr.extend(re.compile(date_regex).findall(match))
                elif re.compile(date_regex_check2).search(match):
                    if re.compile("(?:[Nn]gày)|(?:[Ll]ần ?\d{1} ?: ?"+date_regex_check2+")").search(match):
                        arr.extend(re.compile(date_regex).findall(match))
            return validate_test_dates(arr)
<<<<<<< HEAD
    return None
def extract_positive_place(block_text):
    regex_cach_ly = "(?:(?:chuyển.*)?[Cc][Áá][Cc][Hh] [Ll][Yy].*(?:do))"
    regex = "(?:[Pp]hong [Tt][oỏ][aả])|(?:[Dd]ương tính)|(?:[Tt]heo [Dd]iện)|(?:DƯƠNG TÍNH)|"+regex_cach_ly+"|(?:[Tt]iếp [Xx]úc (?:[Gg]ần)?)"
    regex = re.compile(regex)
    global da_cach_ly
    if regex.search(block_text) and da_cach_ly is False:
        print('Nguon lay',block_text)
        if re.compile("(?:[Tt]iếp [Xx]úc (?:[Gg]ần)?)|(?:[Tt]rong khu cách ly)|(?:F1)|(?:F0)|"+regex_cach_ly).search(block_text):
            print('k')
            return 'Cách ly'
        elif re.compile("[Pp]hong [Tt][oỏ][aả]").search(block_text):
            # print(re.compile("(?:[Pp]hong [Tt][oỏ][aả])").findall(block_text))
            print('e')
            if re.compile("(?:[Gg]ần) (?:(?:(?:[a-z"+VN_regex_norm+"]+) ){1,4})(?:[Pp]hong [Tt][oỏ][aả])").search(block_text) is None:
                return 'Cách ly'
            elif re.compile("trong (?:(?:(?:[a-z"+VN_regex_norm+"]+) ){1,4})(?:[Pp]hong [Tt][oỏ][aả])").search(block_text):
                return 'Cách ly'
        elif re.compile(BN_regex+"|"+BN_regex2+"|(?:xử lý theo quy trình chống dịch)").search(block_text):
            print('d')
            return 'Cách ly'
        else:
            return 'Cộng đồng'
=======
>>>>>>> 3a3ea677d45677e2a850c55fa547dfce3f852eba
    return None
def validate_test_dates(arr):
    valid_arr = []
    for d in arr:
        try:
            if(len(d)<= 2):
                time = datetime.strptime(d,'%d')
                valid_arr.append(time.strftime('%d'))
            elif (len(d)<=5):
                time = datetime.strptime(d, '%d/%m')
                valid_arr.append(time.strftime('%d/%m') + '/2021')
                # return time + datetime.timedelta(days=1)
            else:
                time = datetime.strptime(d, '%d/%m/%Y')
                valid_arr.append(time.strftime('%d/%m/%Y'))
        except:
            pass
    return valid_arr

def extract_publish_date(directory_path):
    dir_name = directory_path[directory_path.rindex('/')+1:]
    publish_date = dir_name.replace('-', '/')
    return publish_date, dir_name

def extract_relative_filepath(file_path):
    file_name = file_path[file_path.rindex('/')+1:]
    
def extract_patient_code_from_filepath(file_path):
    file_name = file_path[file_path.rindex('/')+1:]
    matches = re.compile("(?:BN ?\d+)|(?:Bn ?\d+)").findall(file_name)
    
    if len(matches) > 0:
        patient_code = matches[0]
    else:
        patient_code = None
    return patient_code

def is_valid_patient_code(patient_code):
    if not patient_code.lower().startswith('bn'):
        return False 
    
    for c in patient_code[2:]:
        if c == ' ':
            continue
        if int(c) > 0:
            return True
    return False

def infer_age_info(birthyear):
    age = today_date.year - birthyear
    age_group = ''
    if age < 18:
        return age, '1 (< 18)'
    elif age <= 40:
        return age, '2 (18-40)'
    elif age <= 60:
        return age, '3 (41-60)'
    else:
        return age, '4 (>60)'

def convert_to_report_format(patient_info_json):
    mapping_keys = {
        'publish_date': 'Ngày công bố',
        'patient_code': 'MCB',
        'name': 'Họ và tên',
        'CMND': 'CMND',
        'birthyear': 'Năm sinh',
        'gender': 'Giới',
        'job': 'Nghề nghiệp',
        'job_loc': 'Nơi làm việc/học tập',
        'street': 'Thôn, xóm, đường(thường trú)',
        'ward': 'Xã/Phường(thường trú)',
        'district': 'Quận/Huyện(thường trú)',
        'provine': 'Tỉnh/TP(thường trú)',
        'phone': 'Số điện thoại [bệnh nhân]',
        'test_dates': 'Ngày lấy mẫu',
        'positive_date': 'Ngày xét nghiệm (+)',
        'positve_case_contact': 'Ca F0 liên quan',
        'epidemiology': 'Tóm tắt dịch tễ/ Ghi chú'
    }

    report_json = {}
    for key in patient_info_json:
        if key in mapping_keys:
            if key == 'birthyear' and patient_info_json[key] != '':
                byear = int(patient_info_json[key])
                report_json[mapping_keys[key]] = byear
                report_json['Tuổi'], report_json['Nhóm tuổi'] = infer_age_info(byear)
            else:
                report_json[mapping_keys[key]] = patient_info_json[key]
        else:
            report_json[key] = patient_info_json[key]
    return report_json

def export_to_excel(patient_infos, publish_date, ofile_path):
    df = pd.DataFrame.from_dict(patient_infos, orient='columns') 
    df['Tỉnh/TP báo cáo ca bệnh'] = 'TP HCM'
    # ext_cols = ['Địa chỉ/quê quán', 'Khởi phát', 'Ngày khởi phát', 'Triệu chứng khởi phát', 'Bệnh viện điều trị', 'Nơi điều trị', 'Bệnh nền', 'Tên bệnh nền','Nơi nghi ngờ nhiễm bệnh', 'Nguồn nghi nhiễm', 'Tiếp xúc với ca dương tính', 'Mối quan hệ với ca dương tính', 'Khu công nghiệp', 'Phân loại tiếp xúc', 'Loại ca bệnh', 'Cách phát hiện ca bệnh', 'Ngày tiếp xúc đầu tiên', 'Ngày tiếp xúc cuối cùng', 'Ngày tử vong/hoàn thành điều trị', 'Số điện thoại (có chú thích)', 'Số điện thoại [bệnh nhân]', 'Số điện thoại [bệnh nhân hoặc người nhà', 'CT Value (lần 1)', 'CT Value (lần 2)',	'CT Value (lần 3)', 'CT Value (lần 4)', 'Ngày CT', 'Ngày dịch tễ', 'Nguồn nghi nhiễm', 	'Đối tượng lấy mẫu', 'Tóm tắt dịch tễ/ Ghi chú', 'Gọi điện phỏng vấn thêm F0', 	'TKNC', 'TKND',	'TKSK',	'SĐT theo TKYT', 'Ngày KB', 'Triệu chứng']
    # for col in ext_cols:
    #     df[col] = ''
    
    cols = [
        'Ngày công bố', 'MCB', 'Họ và tên', 'Năm sinh', 'Tuổi', 'Nhóm tuổi',
        'Giới', 'Nghề nghiệp', 'Nơi làm việc/học tập', 'Thôn, xóm, đường(thường trú)',
        'Xã/Phường(thường trú)', 'Quận/Huyện(thường trú)', 'Tỉnh/TP(thường trú)',
        'Tỉnh/TP báo cáo ca bệnh', 'Ngày lấy mẫu', 'Ngày xét nghiệm (+)', 
        'Số điện thoại [bệnh nhân]', 'Tóm tắt dịch tễ/ Ghi chú', 'CMND', 'Ca F0 liên quan'
    ]
    
    print('### Save data to ', ofile_path)
    df[cols].to_excel(ofile_path, sheet_name=publish_date, index=False)
    
# Multiple patients
def extract_multiple_patients(file_path):
    document = Document(file_path)
    
    lines = [paragraph.text.strip() for paragraph in document.paragraphs]
    raw_text = '\n'.join(lines)

    text_blocks = split_normal_multiple(raw_text)
    
    patient_infos = []
    for text_block in text_blocks:
        basic_info = extract_basic_info(text_block)
        location_info = extract_location_info(text_block)
        epi_info = extract_epidemiological_info(text_block)

        patient_info = basic_info.copy()
        patient_info.update(location_info)
        patient_info.update(epi_info)
        patient_infos.append(patient_info)
        
    return patient_infos

def remove_last_line(s):
    return s[:s.rfind('\n')]

def split_normal_multiple(raw_text):
    """Return a list contains infomation for each patient in the word document of normal type"""
    raw_text_lower = raw_text.lower()
    anchor = "thông tin ca bệnh"
    pos = [m.start() for m in re.finditer(anchor, raw_text_lower)]
    pos.append(len(raw_text_lower))
    i = 0
    splitted_patient_info = []
    while i < len(pos) - 1:
        splitted_patient_info.append(raw_text[pos[i]:pos[i+1]])
        i += 1

    i = 0
    while i < len(pos) -2:
        splitted_patient_info[i]= remove_last_line(splitted_patient_info[i].strip())
        i+=1
    return splitted_patient_info  
    
# ----------------------------------------
# input: directory path
# output: a excel file as required
def extract_patient_infos_from_directory(directory_path):
    doc_classifier = DocumentClassifier()
    doc_classes, doc_sizes = doc_classifier.categorize(directory_path)
    
    file_paths = doc_classes['normal_single']
    publish_date, folder_name = extract_publish_date(directory_path)

    count = 0
    patient_infos = []
    ignored_file_paths = []


    for doc_clazz in doc_classes:

        if doc_clazz != 'normal_single' and doc_clazz != 'normal_multiple' and doc_clazz == 'quick_report' and doc_clazz == 'quick_report2':
            ignored_file_paths.extend(doc_classes[doc_clazz])
            continue
        
        file_paths = doc_classes[doc_clazz]
        print('\n', '*'*100)
    
        for file_path in file_paths:
            try:
                print('-'*100, )
                print('@', file_path)

                if doc_clazz == 'quick_report' or doc_clazz == 'quick_report2':
                    patient_infos_from_file = [bao_cao_nhanh.get_personal_information(file_path)]
                    print('---> quickreport', type(patient_infos_from_file), patient_infos_from_file)
                
                if doc_clazz == 'normal_single':
                    patient_infos_from_file = [extract_single_patient(file_path)]
                else:
                    patient_infos_from_file = extract_multiple_patients(file_path)
                
                for patient_info in patient_infos_from_file:
                    
                    if doc_clazz == 'normal_single':
                        patient_code = extract_patient_code_from_filepath(file_path)
                        if not is_valid_patient_code(patient_info['patient_code']) and patient_code is not None:
                            patient_info['patient_code'] = patient_code
                    
                    patient_info['file_path'] = file_path
                    patient_info['publish_date'] = publish_date
                    
                    patient_info = convert_to_report_format(patient_info)
                    
                    print(patient_info)
                    patient_infos.append(patient_info)
            except:
                count += 1
                print('---> error: ', file_path)
                ignored_file_paths.append(file_path)
                
    print('#Extract', len(patient_infos), '#Error:', count, '#All:', doc_sizes['total'])
    print('Info', doc_sizes)
    
    # Convert to the excel format
    export_to_excel(patient_infos, folder_name, directory_path + '.xlsx')
    
    # Log ignored file paths
    if len(ignored_file_paths) > 0:
        with open(directory_path + "_review.txt", 'w') as out:
            for file_path in ignored_file_paths:
                out.write("{}\n".format(file_path))
     
            
if __name__ == '__main__':
    global args
    args = parser.parse_args()
    print(vars(args))
    
    directory_path = args.dir
    extract_patient_infos_from_directory(directory_path)


