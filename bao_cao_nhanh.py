from typing import final
import docx
import os
import re
import pandas as pd
import json
from collections import OrderedDict
from datetime import datetime
from dateutil import parser
from datetime import timedelta

def docx_to_string(docx_file):
    try:
        document = docx.Document(docx_file)
    except:
        print("Unable to read file")
        return ""
    return "\n".join([paragraph.text for paragraph in document.paragraphs])

# làm sạch thông tin các trường cơ bản sau khi áp dụng regex : họ tên , maBN, giới tính, năm sinh
def clean_personal_information(raw_dict):
    cleaned_dict = {}
    if 'hoTen' in raw_dict:
        tien_tos= ["F0",":","Ông", "Bà","bà"]
        hau_tos=['\n']
        hoten = raw_dict["hoTen"][0][0]
        for tien_to in tien_tos:
            hoten = hoten.replace(tien_to, "")
        for hau_to in hau_tos:
            hoten = hoten.replace(hau_to, "")
        cleaned_dict["hoTen"] = hoten
    else:
        cleaned_dict["hoTen"] = ''

    if 'maBN' in raw_dict:
        cleaned_dict["maBN"] = raw_dict["maBN"][0]
    else:
        cleaned_dict["maBN"] = ''

    if 'namSinh' in raw_dict:
        cleaned_dict["namSinh"] = raw_dict["namSinh"][0][-4:]
    else:
        cleaned_dict["namSinh"] = ''

    if 'gioiTinh' in raw_dict:
        cleaned_dict["gioiTinh"] = raw_dict["gioiTinh"][0][1:][1]
    else:
        cleaned_dict["gioiTinh"] = ''

    if "CMND" in raw_dict and len(raw_dict["CMND"][0][0]) > 8:
        cleaned_dict["CMND"] = re.findall('[0-9]+', raw_dict["CMND"][0][0])[0]
    else: 
        cleaned_dict["CMND"] =""
 
    if 'quocTich' in raw_dict:
        startIndex = raw_dict["quocTich"][0][0].find(":")
        cleaned_dict["quocTich"] = raw_dict["quocTich"][0][0][startIndex+1:startIndex+10]
    else:
        cleaned_dict["quocTich"] = ''

    if 'SDT' in raw_dict:
        cleaned_dict["SDT"] = re.findall('[0-9]+', raw_dict["SDT"][0][0])[0]
    else:
        cleaned_dict["SDT"] = ''

    return cleaned_dict

#  các luận để lấy thông tin
def get_personal_information_detail(text):
    regex_dict = {
        "hoTen" : "(((((F0)?\s?:\s?))([A-ZẮẰẲẴẶĂẤẦẨẪẬÂÁÀÃẢẠĐẾỀỂỄỆÊÉÈẺẼẸÍÌỈĨỊỐỒỔỖỘÔỚỜỞỠỢƠÓÒÕỎỌỨỪỬỮỰƯÚÙỦŨỤÝỲỶỸỴ']+\s?){3,5})|((([A-ZẮẰẲẴẶĂẤẦẨẪẬÂÁÀÃẢẠĐẾỀỂỄỆÊÉÈẺẼẸÍÌỈĨỊỐỒỔỖỘÔỚỜỞỠỢƠÓÒÕỎỌỨỪỬỮỰƯÚÙỦŨỤÝỲỶỸỴ][a-záàảãạăắằẳẵặâấầẩẫậéèẻẽẹêếềểễệóòỏõọôốồổỗộơớờởỡợíìỉĩịúùủũụưứừửữựýỳỷỹỵ]{1,}?)\s){3,5}))",
        "maBN":"(BN\s?[0-9]+)",
        "namSinh":"(inh năm[:]*.? \d{4})",
        "gioiTinh":"((:|,)\s?(nam|nữ|NAM|NỮ|Nam|Nữ))",
        "CMND":"(((nhân dân)|(CCCD)|(CMND))\s?:\s?\d{9,})",
        "quocTich": "(tịch\s?:\s?([a-zắằẳẵặăấầẩẫậâáàãảạđếềểễệêéèẻẽẹíìỉĩịốồổỗộôớờởỡợơóòõỏọứừửữựưúùủũụýỳỷỹỵA-ZẮẰẲẴẶĂẤẦẨẪẬÂÁÀÃẢẠĐẾỀỂỄỆÊÉÈẺẼẸÍÌỈĨỊỐỒỔỖỘÔỚỜỞỠỢƠÓÒÕỎỌỨỪỬỮỰƯÚÙỦŨỤÝỲỶỸỴ\s])+)",
        "SDT":"(((thoại)|(ĐT))\s?:\s?\d{4}.?\d{3}.?\d{3})"
    }
    rtn_dict = {}
    for regex in regex_dict:
        if re.search(re.compile(regex_dict[regex]),text):
            rtn_dict[regex] = [m for m in re.findall(re.compile(regex_dict[regex]),text)]
    return rtn_dict

#-----------------------------------------------------------------------------
# ducminh lấy thông tin về nghề nghiệp , địa chỉ

def extract_sections(document_string, section):
   """"Return the string which contains the content of each section"""
   begin = ""
   end = ""
   if section == 1:
      begin = document_string.find("Thông tin ca bệnh")
      end = document_string.find("Lịch sử đi lại và tiền sử tiếp xúc và triệu chứng lâm sàng")
   elif section == 2:
      begin = document_string.find("Lịch sử đi lại và tiền sử tiếp xúc và triệu chứng lâm sàng")
      end = document_string.find("Các hoạt động đã triển khai")
   elif section == 3:
      begin = document_string.find("Các hoạt động đã triển khai")
      end = len(document_string)
   section_string = document_string[begin:end]
   return section_string

def find_job(section_string):
    """Return patient's job description from patient's info string"""
    pattern = re.compile(r'(?<=Nghề nghiệp:)[^\n]*(?=\n)|(?<=Tên và địa chỉ nơi làm việc:)[^\n]*(?=\n)|(?<=Tên và địa chỉ làm việc)[^\n]*(?=\n)|(?<=Địa chỉ nơi ở và nơi làm)[^\n]*(?=\n)')
    job_search = re.search(pattern, section_string)
    if job_search:
        return job_search.group()

def find_address(document_string):
    """Return patient's address from patient's info stirng"""
    address = ""
    pattern = re.compile(r'(?<=Địa chỉ:)[^\n]+|(?<=Địa chỉ nơi ở:)[^\n]+|(?<=Địa chỉ tạm trú:)[^\n]+|(?<=Địa chỉ nơi ở hiện nay:)[^\n]+|(?<=Địa chỉ nơi)[^\n]+|(?<=Địa chỉ nơi ở và nơi làm)[^\n]+|(?<=Địa chỉ nhà:)[^\n]+')
    if re.search(re.compile(pattern),document_string):
        address = re.findall(pattern, document_string)[0]
    return address

def split_address_normal(address_string):
        address_string = address_string.replace('.', '')
        districts = [ 'quận 11','q11', 'quận 12','q12' ,'quận 10','q10','quận 9','q9', 'quận 4','q4' ,   'quận 6','q6',
            'quận 2','q2','quận 5', 'q5', 'quận 7', 'q7','q3' ,'quận 3',  'quận 1', 'q1' ,
            'cần giờ','củ chi','gò vấp', 'phú nhuận',  'bình thạnh', 'quận 8','q8', 'tân bình', 'nhà bè',  
            'hóc môn', 'bình chánh','thủ đức', 'tân phú','bình tân']
        villages = ['phường 16', 'p16', 'tăng nhơn phú b', 'thủ thiêm', 'linh đông', 'bình trưng tây', 'hoà thạnh', 
            'phường 08','p8', 'thảo điền', 'tân chánh hiệp', 'hiệp thành', 'thạnh an', 'vĩnh lộc a', 
            'bình thuận', 'phú xuân', 'tân an hội', 'nhơn đức', 'phường 11','p11' ,'long thới', 'hóc môn', 
            'an nhơn tây', 'phước long b', 'trường thọ', 'bình trưng đông', 'an phú đông', 'bình chiểu', 'bình mỹ',
            'cầu kho', 'thái mỹ', 'phạm văn hai', 'hưng long', 'tân hiệp', 'thạnh xuân', 'trung mỹ tây', 'hiệp bình chánh', 
            'phú thuận', 'tân hưng', 'an lạc', 'tân thới hoà', 'thạnh lộc', 'long thạnh mỹ', 'bình hưng hoà b', 'bà điểm', 
            'tân quý tây', 'bình khánh', 'phường 28', 'p28', 'hiệp phước', 'phú trung', 'bình hưng hòa', 
            'đông hưng thuận', 'phường 7','p7', 'nhị bình', 'tân kiểng', 'an lợi đông', 'linh chiểu', 'an khánh', 
            'phú hòa đông', 'phạm ngũ lão', 'bình an', 'tân thạnh tây', 'phước long a', 'bình chánh', 'linh trung', 'bình trị đông', 'bến nghé', 'long bình', 'tân thuận tây', 'tăng nhơn phú a', 'phước thạnh', 
            'tam phú', 'xuân thới thượng', 'phường 12','p12', 'trung lập thượng', 'tân sơn nhì', 'thới tam thôn', 
            'tân tạo a', 'bình trị đông a', 'phường 24','p24', 'phước bình', 'bình trị đông b', 'phước lộc', 
            'tân qúy', 'đông thạnh', 'lê minh xuân', 'tân nhựt', 'tam bình', 'phú hữu', 'phước kiển', 'tây thạnh', 
            'an thới đông', 'phường 10', 'p10', 'đa kao', 'nhà bè', 'hiệp tân', 'phạm văn cội', 'phú thọ hoà', 
            'quy đức', 'an phú tây', 'tân xuân', 'phú mỹ hưng', 'phường 27', 'p27', 'linh xuân', 'phường 26', 'p26', 
            'nguyễn thái bình', 'phước vĩnh an', 'linh tây', 'tân phong', 'bến thành', 'hiệp phú', 
            'bình hưng hoà a', 'phường 21', 'long phước', 'trung lập hạ', 'phường 17', 'p17', 'tân thạnh đông', 
            'cô giang', 'xuân thới đông', 'bình lợi', 'an phú', 'đa phước', 'phú mỹ', 'tân hưng thuận', 'bình thọ', 
            'phường 18','p18', 'tân tạo', 'phước hiệp', 'cần thạnh', 'phường 13', 'p13','tân quy', 'hiệp bình phước', 
            'phường 15', 'phường 05', 'long hòa', 'hòa phú', 'cầu ông lãnh', 'cát lái', 'thới an', 'tân thuận đông', 
            'tân thới hiệp', 'xuân thới sơn', 'thạnh mỹ lợi', 'tân kiên', 'lý nhơn', 'an lạc a', 'trung an', 'phong phú', 
            'bình hưng', 'nguyễn cư trinh', 'phường 3', 'p3', 'phường 6','p6' ,'phường 19','p19','tân túc', 'phú thạnh', 'phường 14','p14', 'tân phú trung', 'tân phú', 'tân thới nhì', 'phường 22', 'p22',
            'tân thông hội', 'sơn kỳ', 'trung chánh', 'tân định', 'tân thới nhất', 'tam thôn hiệp', 'vĩnh lộc b', 
            'phường 25','p25', 'phường 9','p9', 'củ chi', 'tân thành', 'nhuận đức', 'long trường', 'trường thạnh', 
            'phường 1', 'p1', 'phường 2', 'p2', 'phường 4','p4']
        hcm = {'tp hcm', 'tphcm', 'tp hồ chí minh', 'thành phố hồ chí minh'}
        street = ""
        village = ""
        district = ""
        provine = ""
        for i in hcm:
                if i in address_string:
                    provine = "TP HCM"
                    address_string = address_string[:address_string.find(i)]
                    break
        for d in districts:
                if d in address_string:
                    if d == "tân phú":
                        break
                    district = d
                    provine = "TP HCM"
                    address_string = address_string[:address_string.find(d)]
                    break
        for v in villages:
            if v in address_string:
                if v == "tân phú":
                    break
                village = v
                provine = 'TP HCM'
                street= address_string[:address_string.find(v)]
                street = street.replace("xã", "")
                street = street.replace("phường", "")
                street = street.replace(",", "")
                break
        return street, village, district, provine

def extract_patient_info(patient_info_section):
    """"Return a json file which contains patient's information extracted from a word document"""
    # document_string = docx_to_string(docx_file)
    # patient_info_section = extract_sections(document_string, 1)
    #get patient's job's description
    patient_job = find_job(patient_info_section)
    #get patient's address location
    patient_address = find_address(patient_info_section).lower()
    street, village, district, provine = split_address_normal(patient_address)
    output = {
        # "file_name": docx_file,
        "nghe_nghiep":patient_job,
        "dia chi": patient_address,
        "duong/thon/xom": street,
        "xa/phuong": village,
        "quan/huyen": district,
        "tinh/thanhpho": provine
    }
    return output
  

# #------------------------------------------------------------
# # nghia lấy thông tin dịch tễ

entry_dichte = False
VN_regex_cap = "ẮẰẲẴẶĂẤẦẨẪẬÂÁÀÃẢẠĐẾỀỂỄỆÊÉÈẺẼẸÍÌỈĨỊỐỒỔỖỘÔỚỜỞỠỢƠÓÒÕỎỌỨỪỬỮỰƯÚÙỦŨỤÝỲỶỸỴ"
VN_regex_norm = "áàảãạăắằẳẵặâấầẩẫậéèẻẽẹêếềểễệóòỏõọôốồổỗộơớờởỡợíìỉĩịúùủũụưứừửữự"
# date_regex = "[0-9]{1,2}/[0-9]{1,2}(?:\/[0-9]{4})?"
date_regex = "[0-9]{1,2}/[0-1]{0,1}[0-9]{0,1}(?:\/[0-9]{4})?"
prefix_date_regex = '(?:lấy[^.]*?'+date_regex+')|(?:[Ll]ần.*?'+date_regex+')|(?:'+date_regex+'[^\.]*?lấy mẫu)'
BN_regex = "(?:BN ?\d+)|(?:BN (?:(?:[A-Z"+VN_regex_cap+"]{1,})\s?){2,5})|(?:BN (?:(?:[A-Z"+VN_regex_cap+"][a-z"+VN_regex_norm+"]{1,})\s?){2,5})"


def extract_Ngay_duong_tinh(paragraph):
    regex = "(?:kết quả.*?dương tính[^\.]+?"+date_regex+")|(?:"+date_regex+"[^\./]+kết quả.*?dương tính)"
    regex = re.compile(regex,flags=re.I)
    list_match = None
    if regex.search(paragraph.text):
        # list_match = [m for m in regex.findall(paragraph.text)]
        # print(list_match)
        list_match = regex.findall(paragraph.text)
        # print('ngay_duong_tinh',list_match)
        for match in list_match:
            arr = re.compile(date_regex).findall(match)
        # list_match = list(OrderedDict.fromkeys(list_match))
        return arr[-1]
    else:
        regex_ngay_lay_mau = re.compile(prefix_date_regex)
        # regex_check = re.compile('')
        if regex_ngay_lay_mau.search(paragraph.text):
            arr = extract_Ngay_lay_mau(paragraph)
            # print('arr',arr[-1])
            if(len(arr[-1])<= 2):
                time = datetime.strptime(arr[-1],'%d') + timedelta(days=1)
                return time.strftime('%d')
            elif (len(arr[-1])<=5):
                time = datetime.strptime(arr[-1], '%d/%m') + timedelta(days=1)
                return time.strftime('%d/%m')
                # return time + datetime.timedelta(days=1)
            else:
                time = datetime.strptime(arr[-1], '%d/%m/%Y') + timedelta(days=1)
                return time.strftime('%d/%m/%Y')

    return list_match

def extract_Dich_te(paragraph):
    regex = "[Dd]ịch [Tt]ễ:?.*"
    regex = re.compile(regex)
    global entry_dichte
    if (regex.search(paragraph.text) != None )| entry_dichte:
        if re.compile('\n').search(paragraph.text):
            # print("co dau xuong dong")
            if re.compile('[+]').search(paragraph.text):
                entry_dichte = True
                if entry_dichte:
                    return paragraph.text
                else:
                    return None
            else:
                entry_dichte = False
                return paragraph.text
        else:
            if(paragraph.text.find(':')):
                iter = paragraph.text.find(':')
                return paragraph.text[iter+1:].strip()
    return None

def extract_Ngay_lay_mau(paragraph):
    # regex = "([Dd]ương tính)"
    regex = re.compile(prefix_date_regex)
    arr = []
    if regex.search(paragraph.text):
        # regex = re.compile(prefix_date_regex)
        list_match = regex.findall(paragraph.text)
        # print('ngay_lay_mau',list_match)
        for match in list_match:
            arr.extend(re.compile(date_regex).findall(match))
        return arr
    return None

def extract_Tiep_xuc_ca_duong_tinh(paragraph):
    regex = "(?:[Dd]ương tính)|(?:[Tt]heo [Dd]iện)"
    # ([Tt]iếp xúc)
    regex = re.compile(regex)
    if regex.search(paragraph.text):
        list_match = re.compile(BN_regex).findall(paragraph.text)
        # list_match = list(OrderedDict.fromkeys(list_match))
        # print (paragraph.text)
        # print('Tiep xuc',list_match)
        if len(list_match) == 0:
            return ['Chua ro nguon lay']
        else:
            return list_match
    return None

def single_patient(document):
    Ngay_lay_mau = []
    Ngay_xet_nghiem_duong_tinh = ''
    Dich_te = []
    Tiep_xuc_ca_duong_tinh = []
    i = 0
    Thong_tin_ca_benh = []
    for paragraph in document.paragraphs:
        if 'Thông tin ca bệnh' in paragraph.text:
            i += 1
        if i > 0:
            # print(paragraph.text)
            Thong_tin_ca_benh.append(paragraph)
        if 'Lịch sử đi lại và tiền sử' in paragraph.text:
            i = 0
            break
    # print(Thong_tin_ca_benh)
    for paragraph in Thong_tin_ca_benh:
        if extract_Ngay_duong_tinh(paragraph) != None:
            # print(extract_Ngay_duong_tinh(paragraph))
            Ngay_xet_nghiem_duong_tinh = extract_Ngay_duong_tinh(paragraph)
        if extract_Dich_te(paragraph) != None:
            # print(extract_Dich_te(paragraph))
            Dich_te.append(extract_Dich_te(paragraph))
        if extract_Ngay_lay_mau(paragraph) != None:
            # print(extract_Ngay_lay_mau(paragraph))
            Ngay_lay_mau.extend(extract_Ngay_lay_mau(paragraph))
        if extract_Tiep_xuc_ca_duong_tinh(paragraph) != None:
            # print(extract_Tiep_xuc_ca_duong_tinh(paragraph))
            Tiep_xuc_ca_duong_tinh = extract_Tiep_xuc_ca_duong_tinh(paragraph)
    return {'Ngay_xet_nghiem_duong_tinh':Ngay_xet_nghiem_duong_tinh,
              'Dich_te':Dich_te,
              'Ngay_lay_mau':Ngay_lay_mau,
            'Tiep_xuc_ca_duong_tinh':Tiep_xuc_ca_duong_tinh
              }



# ------------------------------------------------------------------------------------------------------------
#  thông tin cơ bản và địa chỉ
def get_personal_information(file_paths):
    file_paths = [x.replace('\n', '') for x in file_paths]
    count_number_of_patients = 0
    # tong_truong_bat_duoc la 1 dict chưa thông tin xem có tổng bao nhiêu trường bắt được
    tong_truong_bat_duoc={} # tổng của từng trường bắt được trong nhiều file
    thong_tin_ca_nhan = ["hoTen","maBN","namSinh","gioiTinh","CMND","quocTich","SDT"]
    thong_tin_dia_chi = ["nghe_nghiep","dia chi","duong/thon/xom","xa/phuong","quan/huyen","tinh/thanhpho"]

    for index in thong_tin_ca_nhan:
        tong_truong_bat_duoc[index]=0
    for index in thong_tin_dia_chi:
        tong_truong_bat_duoc[index]=0

    for file_path in file_paths:
        count_number_of_patients +=1
        print(count_number_of_patients)

        #  phần này của nghĩa tích hợp để lấy các thông tin về dịch tễ
        # try:
        #     document = docx.Document(file_path)
        #     xet_nghiem_info = single_patient(document)
        # except:
        #     print("no xet_nghiem_info")
        # print(xet_nghiem_info)


        text = docx_to_string(file_path)

        # thông tin các trường: nghề nghiệp, địa chỉ
        address_info = extract_patient_info(text)
        # thông tin các trường: họ tên , năm sinh, giới tính, quốc tịch, cmnd, sdt
        personal_information = get_personal_information_detail(text)
        cleaned_personal_information_dict = clean_personal_information(personal_information)

        #  tổng hợp lại các trường vào 1 dict
        final_info = dict(cleaned_personal_information_dict, **address_info)
        print(final_info)
        print()
        print(file_path) # in ra đường dẫn của file

        # thống kê xem có những trường nào bắt được và trường nào không bắt được
        cac_truong_khong_bat_duoc = []

        for index in cleaned_personal_information_dict:
            if cleaned_personal_information_dict[index] == "":
                cac_truong_khong_bat_duoc.append(index)
            else:
                tong_truong_bat_duoc[index] +=1

        for index in address_info:
            if address_info[index] =="":
                cac_truong_khong_bat_duoc.append(index)
            else:
                tong_truong_bat_duoc[index] +=1

        print("cac truong khong bat duoc: ", cac_truong_khong_bat_duoc)
        print()
        print()

    length_of_total_file = len(file_paths)
    for index in tong_truong_bat_duoc:
        print(index, ":", tong_truong_bat_duoc[index],'/' ,length_of_total_file)

if __name__ == '__main__':
    #đường dẫn các file báo cáo nhanh
    quick_report =r'C:\Users\TRANCONGMINH\Covid19-IE\splitted_files_minh\quick_report.txt'
    quick_report2 = r'C:\Users\TRANCONGMINH\Covid19-IE\splitted_files_minh\quick_report2.txt'

    with open(quick_report,'r', encoding="utf8") as f:
        lines = f.readlines()
    get_personal_information(lines)