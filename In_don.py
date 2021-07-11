# -*- coding: utf-8 -*-
import pandas as pd
import os
import json
import re
from docx import Document
from collections import OrderedDict
from datetime import datetime
from datetime import timedelta


#path docx file
# file_path = '/Users/user/Downloads/BÁO CÁO FILE WORD/BC 24 QUẬN HUYỆN TỪ 1-7/QUẬN TÂN BÌNH/BN000_LÊ THỊ BÍCH TRÂM_30062021.docx'
# file_path = '/Users/user/Downloads/BÁO CÁO FILE WORD/BN0000__HƯNG YÊN_TRẦN VĂN TĂNG_030721.docx'
# file_path = '/Users/user/Downloads/BÁO CÁO FILE WORD/BC CHUỖI CHỢ BÌNH ĐIỀN - NHÓM 2+3/BN0000_ LÊ LÂM THỌ_ĐINH THỊ TRIỀU_24062021.docx'
file_path = '/Users/user/Downloads/covid_path_split_files/arr_path_1.txt'
# file_path = '/Users/user/Downloads/BÁO CÁO FILE WORD/BC 24 QUẬN HUYỆN TỪ 1-7/HUYỆN HÓC MÔN/BN0000_HM_LÊ THỊ HOÀNG_030721.docx'
# file_path = '/Users/user/Downloads/BÁO CÁO FILE WORD/ BC CHUỖI VỰA VE CHAI/BN0000_ĐẶNG NGỌC PHƯƠNG_220621_NHÓM 4.docx'
# file_path = '/Users/user/Downloads/BÁO CÁO FILE WORD/BC CHUỖI CHƯA XÁC ĐỊNH/BN00000_NGUYÊN THIÊN LỘC_26062021_BẰNG_N3.docx'
#multi
# file_path = '/Users/user/Downloads/BÁO CÁO FILE WORD/ BC CHUỖI VỰA VE CHAI/BN0000_BN0000_01_LƯƠNG THỊ THANH THUÝ_NGUYỄN THỊ LỆ THUỶ_01072021.docx'


#create document object
# document = Document(file_path)

entry_dichte = False
entry_dichte2 = False
da_cach_ly = False
VN_regex_cap = "ẮẰẲẴẶĂẤẦẨẪẬÂÁÀÃẢẠĐẾỀỂỄỆÊÉÈẺẼẸÍÌỈĨỊỐỒỔỖỘÔỚỜỞỠỢƠÓÒÕỎỌỨỪỬỮỰƯÚÙỦŨỤÝỲỶỸỴ"
VN_regex_norm = "áàảãạăắằẳẵặâấầẩẫậéèẻẽẹêếềểễệóòỏõọôốồổỗộơớờởỡợíìỉĩịúùủũụưứừửữựýỳỹỷyđ"
# date_regex = "[0-9]{1,2}/[0-9]{1,2}(?:\/[0-9]{4})?"
date_regex = "[0-9]{1,2}/[0-1]{0,1}[0-9]{0,1}(?:\/[0-9]{4})?"
date_regex_check1 = "[0-9]{1,2}/[0-1]{0,1}[0-9]{0,1}/[0-9]{4}"
date_regex_check2 = "[0-3]{0,1}[0-9]{0,1}/[0-1]{0,1}[0-9]{0,1}"
# date_regex = "[0-9]{1,2}/[0-1]{0,1}[0-9]{0,1}/[0-9]{4}"
prefix_date_regex = '(?:lấy[^.]*?'+date_regex+')|(?:[Ll]ần.*?'+date_regex+')|(?:'+date_regex+'[^\.)]*?lấy mẫu)|(?:[Ll][0-9].*?'+date_regex+')'
prefix_date_regex2 = ''
BN_regex = "(?:BN[ _]?\d+)|(?:BN[ _]?(?:(?:[A-Z"+VN_regex_cap+"]{1,})\s?){2,5})|(?:BN[ _]?(?:(?:[A-Z"+VN_regex_cap+"][a-z"+VN_regex_norm+"]{1,})\s?){2,5})|(?:[Bb]ệnh nhân ?(?:(?:[A-Z"+VN_regex_cap+"]{1,}?\s)){2,5})"
BN_regex2 = "(?:F0[ _]?(?:(?:[A-Z"+VN_regex_cap+"]{1,})\s?){2,5})|(?:F0[ _]?(?:(?:[A-Z"+VN_regex_cap+"][a-z"+VN_regex_norm+"]{1,})\s?){2,5})"

def docx_to_string(docx_file):
    try:
        # print(docx_file)
        document = Document(docx_file)
    except:
        print("Unable to read file")
        return ""
    return "\n".join([paragraph.text for paragraph in document.paragraphs])
def extract_sections(document_string, section):
   """"Return the string which contains the content of each section"""
   begin = ""
   end = ""
   if section == 1:
      begin = document_string.find("Thông tin ca bệnh")
      end = document_string.find("Lịch sử đi lại và tiền sử tiếp xúc và triệu chứng lâm sàng")
   elif section == 2:
      begin = document_string.find("Lịch sử đi lại và tiền sử tiếp xúc và triệu chứng lâm sàng")
      end = document_string.find("Các hoạt động đã triển khai")
   elif section == 3:
      begin = document_string.find("Các hoạt động đã triển khai")
      end = len(document_string)
   section_string = document_string[begin:end]
   return section_string


def extract_Ngay_duong_tinh(document_string):
    regex = "(?:kết quả.*?dương tính[^\.]+?"+date_regex+")|(?:"+date_regex+"[^\./]+kết quả.*?dương tính)"
    regex = re.compile(regex,flags=re.I)
    list_match = None
    arr = []
    if entry_dichte:
        return list_match
    else:
        if regex.search(document_string):
            # print (document_string)
            list_match = regex.findall(document_string)
            print('ngay_duong_tinh',list_match)
            for match in list_match:
                arr = re.compile(date_regex).findall(match)
            # list_match = list(OrderedDict.fromkeys(list_match))
            if len(arr) > 0:
                return arr[-1]
            else:
                return list_match
        else:
            regex_ngay_lay_mau = re.compile(prefix_date_regex)
            if regex_ngay_lay_mau.search(document_string):
                # print(document_string)
                arr = extract_Ngay_lay_mau(document_string)
                if len(arr) > 0:
                    print('arr',arr[-1])
                    if(len(arr[-1])<= 2):
                        time = datetime.strptime(arr[-1],'%d') + timedelta(days=1)
                        return time.strftime('%d')
                    elif (len(arr[-1])<=5):
                        time = datetime.strptime(arr[-1], '%d/%m') + timedelta(days=1)
                        return time.strftime('%d/%m')
                    else:
                        time = datetime.strptime(arr[-1], '%d/%m/%Y') + timedelta(days=1)
                        return time.strftime('%d/%m/%Y')
                else:
                    return list_match
    return list_match

def extract_Dich_te(document_string):
    regex = "[Dd]ịch [Tt]ễ:?.*"
    regex = re.compile(regex)
    global entry_dichte
    global entry_dichte2
    # print('entry',entry_dichte)
    if (regex.search(document_string) != None ) or entry_dichte:
        if len(document_string[document_string.find(':')+1:].strip()) == 0 or entry_dichte:
            print("co dau xuong dong")
            entry_dichte = True
            if re.compile('[+]').search(document_string) and entry_dichte:
                entry_dichte2 = True
                if entry_dichte:
                    return document_string
                else:
                    return None
            else:
                # print
                # entry_dichte = False
                print(entry_dichte2)
                if entry_dichte2:
                    entry_dichte2 = False
                    entry_dichte = False
                if entry_dichte and entry_dichte2 is False:
                    if regex.search(document_string) is None:
                        entry_dichte = False
                        return document_string
                    else:
                        entry_dichte = True
                        return None
        else:
            if(document_string.find(':')):
                entry_dichte = False
                iter = document_string.find(':')
                return document_string[iter+1:].strip()
    return None
def extract_Ngay_lay_mau(document_string):
    # regex = "([Dd]ương tính)"
    global entry_dichte
    regex = re.compile(prefix_date_regex)
    arr = []
    # print('entry',entry_dichte)
    if entry_dichte:
        print('Đang xét dịch tễ')
    else:
        if regex.search(document_string):
            # print (document_string)
            # regex = re.compile(prefix_date_regex)
            list_match = regex.findall(document_string)
            print('ngay_lay_mau',list_match)
            for match in list_match:
                if re.compile(date_regex_check1).search(match):
                    arr.extend(re.compile(date_regex).findall(match))
                elif re.compile(date_regex_check2).search(match):
                    if re.compile("[Nn]gày").search(match):
                        arr.extend(re.compile(date_regex).findall(match))
            return arr
    return None
def extract_Tiep_xuc_ca_duong_tinh(document_string):
    regex = "(?:[Dd]ương tính)|(?:[Tt]heo [Dd]iện)|(?:[Tt]iếp [Xx]úc (?:[Gg]ần)?)|(?:[Ll]iên quan)"
    # ([Tt]iếp xúc)
    regex = re.compile(regex)
    if regex.search(document_string):
        if re.compile("[Bb]ệnh ?[Nn]hân:").search(document_string) is None:
            list_match = re.compile(BN_regex+"|"+BN_regex2).findall(document_string)
            # list_match = list(OrderedDict.fromkeys(list_match))
            print('Tiep xuc',list_match)
            if len(list_match) == 0:
                return None
            else:
                return list_match
    return None
def extract_Nguon_lay_nhiem(document_string):
    regex_cach_ly = "(?:(?:chuyển.*)?[Cc][Áá][Cc][Hh] [Ll][Yy].*(?:do))"
    regex = "(?:[Pp]hong [Tt][oỏ][aả])|(?:[Dd]ương tính)|(?:[Tt]heo [Dd]iện)|(?:DƯƠNG TÍNH)|"+regex_cach_ly
    regex = re.compile(regex)
    global da_cach_ly
    if regex.search(document_string) and da_cach_ly is False:
        print('Nguon lay',document_string)
        if re.compile("(?:[Tt]iếp [Xx]úc (?:[Gg]ần)?)|(?:[Tt]rong khu cách ly)|(?:F1)|(?:F0)|"+regex_cach_ly).search(document_string):
            print('k')
            return 'Cách ly'
        elif re.compile("[Pp]hong [Tt][oỏ][aả]").search(document_string):
            # print(re.compile("(?:[Pp]hong [Tt][oỏ][aả])").findall(document_string))
            print('e')
            if re.compile("(?:[Gg]ần) (?:(?:(?:[a-z"+VN_regex_norm+"]+) ){1,4})(?:[Pp]hong [Tt][oỏ][aả])").search(document_string) is None:
                return 'Cách ly'
            elif re.compile("trong (?:(?:(?:[a-z"+VN_regex_norm+"]+) ){1,4})(?:[Pp]hong [Tt][oỏ][aả])").search(document_string):
                return 'Cách ly'
        elif re.compile(BN_regex+"|"+BN_regex2).search(document_string):
            print('d')
            return 'Cách ly'
    return None
def single_patient(document_string):
    Ngay_lay_mau = []
    Ngay_xet_nghiem_duong_tinh = ''
    Dich_te = []
    Tiep_xuc_ca_duong_tinh = []
    i = 0
    Nguon_lay_nhiem = ''
    global da_cach_ly

    print(extract_sections(document_string, 1))
    Thong_tin_ca_benh = extract_sections(document_string, 1).split('\n')
    # print(Thong_tin_ca_benh)
    for paragraph in Thong_tin_ca_benh:
        # print(paragraph)
        res = extract_Dich_te(paragraph)
        if res != None:
            Dich_te.append(res)

        res = extract_Ngay_duong_tinh(paragraph)
        if res != None:
            Ngay_xet_nghiem_duong_tinh = res

        res = extract_Ngay_lay_mau(paragraph)
        if res != None:
            Ngay_lay_mau.extend(res)

        res = extract_Tiep_xuc_ca_duong_tinh(paragraph)
        if res != None:
            Tiep_xuc_ca_duong_tinh = res

        res = extract_Nguon_lay_nhiem(paragraph)
        if (res == 'Cách ly'):
            da_cach_ly = True
        if res != None:
            Nguon_lay_nhiem = res
# handle data:
    da_cach_ly = False
    if Nguon_lay_nhiem == '':
        Nguon_lay_nhiem = 'Cộng Đồng'

    Ngay_lay_mau = list(OrderedDict.fromkeys(Ngay_lay_mau))

    if len(Ngay_lay_mau) == 0 and len(Ngay_xet_nghiem_duong_tinh) != 0:
        Ngay_lay_mau.append(Ngay_xet_nghiem_duong_tinh)
    # if len(Dich_te) > 0:
    return {'Ngay_xet_nghiem_duong_tinh':Ngay_xet_nghiem_duong_tinh,
              'Dich_te':Dich_te,
              'Ngay_lay_mau':Ngay_lay_mau,
            'Tiep_xuc_ca_duong_tinh':Tiep_xuc_ca_duong_tinh,
            'Nguon lay':Nguon_lay_nhiem
              }


# # run multiple single docx
i = 1;
output = []
with open(file_path, 'r', encoding= 'utf-8') as f:
    for line in f:
        # print(os.path.realpath('./'))path
        # print(line)
        # path = '/Users/user/Downloads/BÁO CÁO FILE WORD/'+line[15:-1]
        path = "/Users/user/Downloads/baocao_covid/BN0000__HƯNG YÊN_TRẦN VĂN TĂNG_030721.docx"
        print('here',path)
        document_string = docx_to_string(path)
        res = single_patient(document_string)
        print('\n',res,'\n')
        print(i)
        i+=1
        if i == (3 + 2):
            break

# test single docx
# document = Document(file_path)
# a = docx_to_string(file_path)
# print('\n',single_patient(a),'\n')




# regex = re.compile(BN_regex)
# list_match = regex.findall('Ngày lấy mẫu xét nghiệm: lấy mẫu lần 1 vào trưa ngày 21/06/2021 tại KCL là  trường học quận 1 (BN không biết tên KCL và địa chỉ) do tiếp xúc gần với BN12399 Lê Thị Ngọc Hương.')
# print(list_match)

# def extract_Dich_te(paragraph):
#     regex = "[Dd]ịch [Tt]ễ:?.*"
#     regex = re.compile(regex)
#     if regex.search(paragraph):
#         if re.compile('\n').search(paragraph):
#             return None
#         else:
#             if(paragraph.find(':')):
#                 iter = paragraph.find(':')
#                 return paragraph[iter:].strip()
#     return None
#
# def multi_patient(document):
#     i = 0
#     documents = []
#     doc = ''
#     for paragraph in document.paragraphs:
#         # print(paragraph)
#         if 'Thông tin ca bệnh' in paragraph:
#             documents.append(doc)
#             doc = Document(docx=None)
#             print('changeeeee')
#             continue
#         # print(type(doc))
#         if 'Các hoạt động đã triển khai' in paragraph:
#             documents.append(doc)
#             break;
#         if type(doc) != str:
#             doc.add_paragraph(paragraph)
#
#     for d in range(1,len(documents)):
#         print(len(documents))
#         print(single_patient(documents[d]))
    # print(documents[1].paragraphs)
    # for p in documents[4].paragraphs:
    #     print(p)
# multi_patient(document)


